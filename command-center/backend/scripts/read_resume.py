#!/usr/bin/env python3
"""
Dump the structure of a resume .docx template.

Shows every paragraph with:
  - Index, style name
  - Run-level detail: bold, italic, font, size, text
  - <w:br/> line break positions
  - Total character count per paragraph

Usage:
    python3 read_resume.py <path_to_docx>
"""

import sys
from docx import Document
from docx.oxml.ns import qn


def pt(p):
    """Get full text of a paragraph, representing <w:br/> as newlines."""
    parts = []
    for child in p._element.iter():
        if child.tag == qn('w:t'):
            parts.append(child.text or '')
        elif child.tag == qn('w:br'):
            parts.append('\n')
    return ''.join(parts)


def dump_paragraph(i, p):
    """Print detailed info about a paragraph."""
    text = pt(p)
    if not text.strip():
        return  # Skip empty paragraphs

    style = p.style.name if p.style else "None"
    br_count = len(p._element.findall('.//' + qn('w:br')))

    print(f"\n{'='*80}")
    print(f"PARAGRAPH {i} | style: {style} | <w:br/> count: {br_count} | chars: {len(text)}")
    print(f"{'='*80}")

    # Show run-level detail
    for j, run in enumerate(p.runs):
        rtext = run.text
        if not rtext and not run._element.findall(qn('w:br')):
            continue

        bold = run.bold
        italic = run.italic
        font = run.font.name if run.font and run.font.name else "-"
        size = run.font.size if run.font and run.font.size else "-"

        # Check for <w:br/> in this run
        has_br = len(run._element.findall(qn('w:br'))) > 0

        flags = []
        if bold:
            flags.append("BOLD")
        if italic:
            flags.append("ITALIC")
        if has_br:
            flags.append("<w:br/>")

        flag_str = f" [{', '.join(flags)}]" if flags else ""
        text_preview = repr(rtext) if rtext else "(empty)"

        print(f"  Run {j}: font={font}, size={size}{flag_str}")
        print(f"    Text: {text_preview}")

    # Show the full assembled text with line markers
    if br_count > 0:
        print(f"\n  --- Full text (lines separated by <w:br/>) ---")
        for k, line in enumerate(text.split('\n')):
            print(f"  Line {k}: {line}")
    else:
        print(f"\n  Full text: {text}")


def main():
    if len(sys.argv) < 2:
        print("Usage: python3 read_resume.py <path_to_docx>")
        sys.exit(1)

    docx_path = sys.argv[1]
    doc = Document(docx_path)

    print(f"Resume Template Analysis: {docx_path}")
    print(f"Total paragraphs: {len(doc.paragraphs)}")

    # Also check for tables (some resumes use tables for layout)
    if doc.tables:
        print(f"Tables found: {len(doc.tables)}")
        for t_idx, table in enumerate(doc.tables):
            print(f"\n  Table {t_idx}: {len(table.rows)} rows x {len(table.columns)} cols")
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    if cell_text:
                        print(f"    [{r_idx},{c_idx}]: {cell_text[:80]}{'...' if len(cell_text) > 80 else ''}")

    # Dump all paragraphs
    for i, p in enumerate(doc.paragraphs):
        dump_paragraph(i, p)

    # Summary
    print(f"\n{'='*80}")
    print("SUMMARY")
    print(f"{'='*80}")

    total_br = 0
    sections = []
    for i, p in enumerate(doc.paragraphs):
        text = pt(p).strip()
        br_count = len(p._element.findall('.//' + qn('w:br')))
        total_br += br_count

        # Detect section headers (typically short, bold, all-caps or Title case)
        if text and len(text) < 40:
            all_bold = all(r.bold for r in p.runs if r.text.strip())
            if all_bold:
                sections.append(f"  P{i}: {text}")

    print(f"Total <w:br/> elements across all paragraphs: {total_br}")
    print(f"Detected section headers (short + bold):")
    for s in sections:
        print(s)


if __name__ == "__main__":
    main()
