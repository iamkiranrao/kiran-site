#!/usr/bin/env python3
"""
Post-edit quality checks for a customized resume.

Checks:
  1. Repeated action verbs across bullet points
  2. Bold formatting on bullet text (should be non-bold)
  3. Metric formatting (no space before M/K/%)
  4. Interests section presence (for 2-Page and Detailed)
  5. Bullet length (warn if > 120 chars)

Usage:
    python3 verify_resume.py <docx_path> [--check-interests]
"""

import sys
import re
import argparse
from docx import Document
from docx.oxml.ns import qn


def pt(p):
    """Get full text of a paragraph, representing <w:br/> as newlines."""
    parts = []
    for child in p._element.iter():
        if child.tag == qn('w:t'):
            parts.append(child.text or '')
        elif child.tag == qn('w:br'):
            parts.append('\n')
    return ''.join(parts)


def check_repeated_verbs(doc):
    """Check for repeated action verbs across all bullet points."""
    issues = []
    verb_locations = {}  # verb -> list of (paragraph_index, line_text)

    for i, p in enumerate(doc.paragraphs):
        text = pt(p)
        lines = text.split('\n')

        for line in lines:
            line = line.strip()
            if not line or len(line) < 15:
                continue

            # Extract first word as potential action verb
            words = line.split()
            if not words:
                continue

            verb = words[0].rstrip(',.:;')

            # Skip non-verb-like first words (names, numbers, section headers)
            if verb[0].isdigit() or len(verb) < 3:
                continue
            if verb.isupper() and len(verb) > 3:
                continue  # Likely a section header like "EXPERIENCE"

            verb_lower = verb.lower()
            if verb_lower not in verb_locations:
                verb_locations[verb_lower] = []
            verb_locations[verb_lower].append((i, line[:80]))

    for verb, locations in verb_locations.items():
        if len(locations) > 1:
            issues.append({
                "type": "repeated_verb",
                "verb": verb,
                "count": len(locations),
                "locations": locations
            })

    return issues


def check_bold_bullets(doc):
    """Check for bullet text that is accidentally bold."""
    issues = []

    for i, p in enumerate(doc.paragraphs):
        text = pt(p)
        br_count = len(p._element.findall('.//' + qn('w:br')))

        # Only check paragraphs that look like bullet sections (have <w:br/>)
        if br_count == 0:
            continue

        for j, run in enumerate(p.runs):
            if run.bold and run.text.strip():
                # Check if this looks like bullet text (not a title)
                if len(run.text.strip()) > 30:
                    issues.append({
                        "type": "bold_bullet",
                        "paragraph": i,
                        "run": j,
                        "text": run.text[:60]
                    })

    return issues


def check_metric_formatting(doc):
    """Check for spaces before M/K/% in metrics."""
    issues = []
    # Pattern: number followed by space then M/K/%/B
    bad_pattern = re.compile(r'\d\s+[MKB%](?:\b|$)')

    for i, p in enumerate(doc.paragraphs):
        text = pt(p)
        matches = bad_pattern.findall(text)
        if matches:
            issues.append({
                "type": "metric_spacing",
                "paragraph": i,
                "matches": matches,
                "text": text[:100]
            })

    return issues


def check_bullet_length(doc, max_chars=120):
    """Warn about bullets that might not fit on one line."""
    issues = []

    for i, p in enumerate(doc.paragraphs):
        text = pt(p)
        lines = text.split('\n')

        for line_num, line in enumerate(lines):
            line = line.strip()
            if len(line) > max_chars:
                issues.append({
                    "type": "long_bullet",
                    "paragraph": i,
                    "line": line_num,
                    "length": len(line),
                    "text": line[:80] + "..."
                })

    return issues


def check_interests_section(doc):
    """Check if an Interests section exists."""
    for p in doc.paragraphs:
        text = pt(p).strip().lower()
        if 'interest' in text and len(text) < 30:
            return True
    return False


def main():
    parser = argparse.ArgumentParser(description="Verify resume quality")
    parser.add_argument("docx_path", help="Path to the .docx file")
    parser.add_argument("--check-interests", action="store_true",
                        help="Check for Interests section presence")

    args = parser.parse_args()
    doc = Document(args.docx_path)

    all_issues = []
    warnings = 0
    errors = 0

    print(f"Verifying: {args.docx_path}")
    print(f"{'='*60}")

    # 1. Repeated verbs
    verb_issues = check_repeated_verbs(doc)
    for issue in verb_issues:
        print(f"\n  WARN: Repeated verb '{issue['verb']}' used {issue['count']} times:")
        for loc in issue['locations']:
            print(f"    P{loc[0]}: {loc[1]}")
        warnings += 1
    all_issues.extend(verb_issues)

    # 2. Bold bullets
    bold_issues = check_bold_bullets(doc)
    for issue in bold_issues:
        print(f"\n  ERROR: Bold bullet text at P{issue['paragraph']}, Run {issue['run']}:")
        print(f"    {issue['text']}")
        errors += 1
    all_issues.extend(bold_issues)

    # 3. Metric formatting
    metric_issues = check_metric_formatting(doc)
    for issue in metric_issues:
        print(f"\n  WARN: Metric spacing issue at P{issue['paragraph']}:")
        print(f"    Found: {issue['matches']}")
        warnings += 1
    all_issues.extend(metric_issues)

    # 4. Bullet length
    length_issues = check_bullet_length(doc)
    for issue in length_issues:
        print(f"\n  WARN: Long bullet ({issue['length']} chars) at P{issue['paragraph']}, line {issue['line']}:")
        print(f"    {issue['text']}")
        warnings += 1
    all_issues.extend(length_issues)

    # 5. Interests section
    if args.check_interests:
        has_interests = check_interests_section(doc)
        if not has_interests:
            print(f"\n  ERROR: Interests section not found (required for 2-Page/Detailed)")
            errors += 1

    # Summary
    print(f"\n{'='*60}")
    print(f"Results: {errors} errors, {warnings} warnings")
    if errors == 0 and warnings == 0:
        print("All checks passed!")

    return 0 if errors == 0 else 1


if __name__ == "__main__":
    sys.exit(main())
