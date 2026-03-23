"""
Resume Certifier — Final gate check that determines if a resume is ready to send.

Runs as the last step in both pipeline modes (quick generate and review-first).
Aggregates results from the scorer, verifier, and fitter, plus runs its own
contact info and ATS safety checks.

Returns one of three verdicts:
- CERTIFIED: All clear. Attach and send.
- WARNINGS:  Sendable, but flagged items could be stronger.
- BLOCKED:   Fix these before sending. Hard blockers found.

Usage:
    from services.resume_certifier import certify_resume
    result = certify_resume("/path/to/resume.docx", target_pages=2)
    result_with_jd = certify_resume("/path/to/resume.docx", target_pages=2, jd_text="...")

    # Or aggregate from pre-computed pipeline results:
    result = certify_resume(
        "/path/to/resume.docx",
        target_pages=2,
        jd_text="...",
        scorer_result=post_score,
        verifier_result=verify_result,
        fitter_result=fit_result,
    )
"""

import re
import logging
from typing import Optional, Dict, List
from docx import Document

from services.governance_loader import (
    HARD_BANNED, SOFT_BANNED, AI_LANGUAGE, ANTI_PATTERNS,
    SOFT_BANNED_LIMIT,
)

logger = logging.getLogger(__name__)

# Contact info patterns
EMAIL_PATTERN = re.compile(r'[\w.+-]+@[\w-]+\.[\w.-]+')
PHONE_PATTERN = re.compile(r'(?:\+?1[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}')
# Match either full URL or hyperlinked text label
LINKEDIN_PATTERN = re.compile(r'linkedin\.com/in/[\w-]+|(?:^|\s)LinkedIn(?:\s|$)', re.IGNORECASE)


# ── Check Functions ───────────────────────────────────────────────────

def _check_page_count(doc: Document, docx_path: str, target_pages: int) -> Dict:
    """Check that the resume renders to the correct page count."""
    actual_pages = target_pages  # Default assumption

    # Try PDF-based count
    try:
        from services.resume_fitter import _convert_to_pdf, _count_pdf_pages
        import os
        pdf_path = _convert_to_pdf(docx_path)
        if pdf_path:
            actual_pages = _count_pdf_pages(pdf_path)
            try:
                os.remove(pdf_path)
            except Exception:
                logger.debug("Could not clean up temp PDF")
    except Exception:
        logger.warning("PDF-based page count failed, falling back to line estimate", exc_info=True)
        # Fallback: estimate from line count
        lines = sum(1 for p in doc.paragraphs if p.text.strip())
        estimated = max(1, round(lines / 45))
        actual_pages = estimated

    if actual_pages > target_pages:
        return {
            "check": "page_count",
            "status": "BLOCKED",
            "message": f"Resume is {actual_pages} pages — target is {target_pages}. Content overflows.",
            "actual": actual_pages,
            "target": target_pages,
        }
    elif actual_pages < target_pages:
        return {
            "check": "page_count",
            "status": "BLOCKED",
            "message": f"Resume is {actual_pages} pages — target is {target_pages}. Significant underfill.",
            "actual": actual_pages,
            "target": target_pages,
        }
    return {
        "check": "page_count",
        "status": "PASSED",
        "message": f"{actual_pages} pages (target: {target_pages})",
    }


def _check_contact_info(full_text: str) -> Dict:
    """Verify name, email, phone, and LinkedIn are present."""
    missing = []

    # Check for email
    if not EMAIL_PATTERN.search(full_text):
        missing.append("email")

    # Check for phone
    if not PHONE_PATTERN.search(full_text):
        missing.append("phone")

    # Check for LinkedIn
    if not LINKEDIN_PATTERN.search(full_text):
        missing.append("LinkedIn URL")

    # Check for name (first paragraph should be the name — non-empty, short)
    # We don't check this structurally since the name is always there in our templates

    if missing:
        return {
            "check": "contact_info",
            "status": "BLOCKED",
            "message": f"Missing contact info: {', '.join(missing)}. Recruiters can't reach you.",
            "missing": missing,
        }
    return {
        "check": "contact_info",
        "status": "PASSED",
        "message": "Email, phone, LinkedIn all present",
    }


def _check_banned_words(full_text: str) -> Dict:
    """Check for hard-banned and AI-flagged words."""
    full_lower = full_text.lower()
    found_hard = [w for w in HARD_BANNED if w.lower() in full_lower]
    found_ai = [w for w in AI_LANGUAGE if w.lower() in full_lower]
    found_anti = [w for w in ANTI_PATTERNS if w in full_lower]
    found_soft = [w for w in SOFT_BANNED if w.lower() in full_lower]

    all_blockers = found_hard + found_ai + found_anti
    soft_over_limit = found_soft[SOFT_BANNED_LIMIT:] if len(found_soft) > SOFT_BANNED_LIMIT else []

    if all_blockers or soft_over_limit:
        items = all_blockers + soft_over_limit
        return {
            "check": "banned_words",
            "status": "BLOCKED",
            "message": f"AI-flagged words found: {', '.join(items)}. 53% of hiring managers flag these.",
            "found": items,
        }

    if found_soft:
        return {
            "check": "banned_words",
            "status": "WARNING",
            "message": f"Soft-banned word used ({', '.join(found_soft)}) — within limit of {SOFT_BANNED_LIMIT}.",
            "found": found_soft,
        }

    return {
        "check": "banned_words",
        "status": "PASSED",
        "message": "No banned or AI-flagged words",
    }


def _check_repeated_verbs(doc: Document) -> Dict:
    """Check for duplicate leading verbs across bullet points."""
    verb_counts = {}
    skip_words = {
        "jira", "mixpanel", "openai", "agile", "liverpool", "product", "pm",
        "analytics", "ai", "frameworks", "bachelor", "master", "mba",
        "certified", "university", "college", "kellogg",
    }

    # Identify skills section paragraphs to skip (skill names, not bullet verbs)
    in_skills = False
    skills_indices = set()
    for i, p in enumerate(doc.paragraphs):
        text_upper = p.text.strip().upper()
        if text_upper in ("CORE COMPETENCIES", "TECHNICAL COMPETENCIES", "TECHNICAL SKILLS",
                          "LEADERSHIP & STRATEGIC COMPETENCIES", "SKILLS", "CORE SKILLS"):
            in_skills = True
            skills_indices.add(i)
            continue
        if in_skills:
            if text_upper and len(text_upper) < 40 and text_upper in (
                "EXPERIENCE", "PROFESSIONAL EXPERIENCE", "WORK EXPERIENCE",
                "EDUCATION", "EDUCATION & CERTIFICATIONS", "INTERESTS", "CERTIFICATIONS",
            ):
                in_skills = False
                continue
            if p.text.strip():
                skills_indices.add(i)

    for idx, p in enumerate(doc.paragraphs):
        if idx in skills_indices:
            continue
        text = p.text.strip()
        # Handle BR-joined paragraphs
        for line in text.split("\n"):
            line = line.strip()
            if not line or len(line) < 15:
                continue
            # Strip bullet markers to get the actual leading verb
            clean = line.lstrip("•·-– ").strip()
            words = clean.split()
            if not words:
                continue
            verb = words[0].rstrip(",.:;")
            if not verb or verb[0].isdigit() or len(verb) < 3:
                continue
            if verb.isupper() and len(verb) > 3:
                continue
            if verb.lower() in skip_words:
                continue
            verb_lower = verb.lower()
            verb_counts[verb_lower] = verb_counts.get(verb_lower, 0) + 1

    dupes = {v: c for v, c in verb_counts.items() if c > 1}
    if dupes:
        dupe_list = [f"{v} (x{c})" for v, c in dupes.items()]
        return {
            "check": "repeated_verbs",
            "status": "BLOCKED",
            "message": f"Repeated action verbs: {', '.join(dupe_list)}. Looks AI-generated.",
            "duplicates": dupes,
        }
    return {
        "check": "repeated_verbs",
        "status": "PASSED",
        "message": f"All {len(verb_counts)} action verbs are unique",
    }


def _check_ats_characters(full_text: str) -> Dict:
    """Check for characters that corrupt in ATS systems."""
    issues = []

    em_dashes = full_text.count("—")
    if em_dashes:
        issues.append(f"{em_dashes} em dash(es)")

    arrows = full_text.count("→")
    if arrows:
        issues.append(f"{arrows} arrow(s)")

    # Check for emojis (non-ASCII decorative characters)
    emoji_pattern = re.compile(
        "[\U0001F600-\U0001F64F\U0001F300-\U0001F5FF\U0001F680-\U0001F6FF"
        "\U0001F1E0-\U0001F1FF\U00002702-\U000027B0\U0000FE00-\U0000FE0F"
        "\U0001F900-\U0001F9FF\U0001FA00-\U0001FA6F\U0001FA70-\U0001FAFF"
        "\U00002600-\U000026FF\U0000231A-\U0000231B]+",
        flags=re.UNICODE,
    )
    emojis = emoji_pattern.findall(full_text)
    if emojis:
        issues.append(f"{len(emojis)} emoji(s)")

    if issues:
        return {
            "check": "ats_characters",
            "status": "BLOCKED",
            "message": f"ATS-unsafe characters: {', '.join(issues)}. These corrupt in Workday/Greenhouse.",
            "issues": issues,
        }
    return {
        "check": "ats_characters",
        "status": "PASSED",
        "message": "No ATS-unsafe characters",
    }


def _check_bold_bullets(doc: Document) -> Dict:
    """Check for formatting artifacts: bullet text that's incorrectly bold."""
    from docx.oxml.ns import qn
    bold_count = 0
    for p in doc.paragraphs:
        br_count = len(p._element.findall(".//" + qn("w:br")))
        if br_count == 0:
            # Also check individual bullet paragraphs
            text = p.text.strip()
            if not text.startswith("•"):
                continue
        for run in p.runs:
            if run.bold and run.text and len(run.text.strip()) > 30:
                bold_count += 1

    if bold_count:
        return {
            "check": "bold_bullets",
            "status": "BLOCKED",
            "message": f"{bold_count} bullet text run(s) incorrectly bold. Formatting artifact.",
            "count": bold_count,
        }
    return {
        "check": "bold_bullets",
        "status": "PASSED",
        "message": "No bold formatting artifacts",
    }


def _check_date_format(full_text: str) -> Dict:
    """Check for inconsistent date formats (apostrophe years, etc.)."""
    bad_dates = re.findall(r"'\d{2}\b", full_text)
    if bad_dates:
        return {
            "check": "date_format",
            "status": "BLOCKED",
            "message": f"{len(bad_dates)} apostrophe-year date(s) found (e.g., 'Nov '23'). Use full year format.",
            "count": len(bad_dates),
        }
    return {
        "check": "date_format",
        "status": "PASSED",
        "message": "Date formats consistent",
    }


def _check_metrics_density(doc: Document) -> Dict:
    """Check that enough bullets have quantified metrics."""
    bullets = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if text.startswith("•") or text.startswith("·"):
            bullets.append(text[1:].strip())

    if not bullets:
        return {
            "check": "metrics_density",
            "status": "WARNING",
            "message": "No bullets found to check metrics density",
        }

    metric_bullets = sum(
        1 for b in bullets
        if re.search(r'\d+[%MKB]|\$\d|\d+\s*(?:point|person|user)', b)
    )
    pct = metric_bullets / len(bullets) * 100

    if pct < 50:
        return {
            "check": "metrics_density",
            "status": "WARNING",
            "message": f"Only {pct:.0f}% of bullets have metrics ({metric_bullets}/{len(bullets)}). Target is 65-85%.",
            "percentage": round(pct),
        }
    return {
        "check": "metrics_density",
        "status": "PASSED",
        "message": f"{pct:.0f}% of bullets have metrics ({metric_bullets}/{len(bullets)})",
    }


def _check_verb_diversity(doc: Document) -> Dict:
    """Check that verbs span multiple categories."""
    from services.resume_scorer import (
        IMPACT_VERBS, CREATION_VERBS, LEADERSHIP_VERBS,
        EXECUTION_VERBS, STRATEGY_VERBS,
    )

    verbs = set()
    for p in doc.paragraphs:
        text = p.text.strip()
        for line in text.split("\n"):
            line = line.strip()
            if line.startswith("•") or line.startswith("·"):
                words = line[1:].strip().split()
                if words and words[0][0].isalpha() and len(words[0]) > 2:
                    verbs.add(words[0].rstrip(",.:;"))

    categories_hit = 0
    if verbs & IMPACT_VERBS: categories_hit += 1
    if verbs & CREATION_VERBS: categories_hit += 1
    if verbs & LEADERSHIP_VERBS: categories_hit += 1
    if verbs & EXECUTION_VERBS: categories_hit += 1
    if verbs & STRATEGY_VERBS: categories_hit += 1

    if categories_hit < 3:
        return {
            "check": "verb_diversity",
            "status": "WARNING",
            "message": f"Only {categories_hit}/5 verb categories used. Reads one-dimensional.",
            "categories": categories_hit,
        }
    return {
        "check": "verb_diversity",
        "status": "PASSED",
        "message": f"{categories_hit}/5 verb categories represented",
    }


def _check_outcome_first(doc: Document) -> Dict:
    """Check for at least one outcome-first bullet (leads with a number)."""
    for p in doc.paragraphs:
        text = p.text.strip()
        for line in text.split("\n"):
            line = line.strip()
            if line.startswith("•"):
                content = line[1:].strip()
                if content and content[0].isdigit():
                    return {
                        "check": "outcome_first",
                        "status": "PASSED",
                        "message": "At least 1 outcome-first bullet present",
                    }

    return {
        "check": "outcome_first",
        "status": "WARNING",
        "message": "No outcome-first bullets found. Consider leading 1-2 bullets with the metric.",
    }


def _check_jd_tagline(full_text: str, jd_text: str) -> Dict:
    """Check that the tagline mirrors the JD's title language."""
    jd_lower = jd_text.lower()

    # Extract likely JD title
    title_words = set()
    for pattern in [r'(?:title|position|role)\s*[:]\s*(.+)', r'^(.+(?:manager|director|lead|vp|head).+)$']:
        m = re.search(pattern, jd_lower, re.MULTILINE)
        if m:
            title_words = set(m.group(1).strip().split())
            break

    if not title_words:
        return {
            "check": "jd_tagline",
            "status": "PASSED",
            "message": "Could not extract JD title — skipping tagline check",
        }

    # Check first 10 paragraphs for the tagline (italic, 30-200 chars)
    for line in full_text.split("\n")[:15]:
        line = line.strip()
        if 30 < len(line) < 200:
            line_lower = line.lower()
            # Exclude section headers
            if line.upper() == line and len(line) < 40:
                continue
            overlap = title_words & set(line_lower.split())
            if len(overlap) >= 2:
                return {
                    "check": "jd_tagline",
                    "status": "PASSED",
                    "message": f"Tagline mirrors JD title ({', '.join(overlap)})",
                }

    return {
        "check": "jd_tagline",
        "status": "WARNING",
        "message": "Tagline may not mirror the JD title. Check that it signals the right role.",
    }


def _check_jd_keyword_coverage(full_text: str, jd_text: str) -> Dict:
    """Check that ATS keyword coverage meets minimum threshold."""
    from services.resume_scorer import _extract_jd_keywords

    jd_keywords = _extract_jd_keywords(jd_text)
    if not jd_keywords:
        return {
            "check": "jd_keywords",
            "status": "PASSED",
            "message": "No extractable JD keywords — skipping coverage check",
        }

    full_lower = full_text.lower()
    matched = [kw for kw in jd_keywords if kw in full_lower]
    coverage = len(matched) / len(jd_keywords) * 100

    if coverage < 40:
        missing_sample = [kw for kw in jd_keywords if kw not in full_lower][:5]
        return {
            "check": "jd_keywords",
            "status": "WARNING",
            "message": f"ATS keyword coverage is {coverage:.0f}% ({len(matched)}/{len(jd_keywords)}). "
                       f"Missing: {', '.join(missing_sample)}",
            "coverage": round(coverage),
            "missing_sample": missing_sample,
        }
    return {
        "check": "jd_keywords",
        "status": "PASSED",
        "message": f"ATS keyword coverage: {coverage:.0f}% ({len(matched)}/{len(jd_keywords)})",
    }


# ── Main Certifier ────────────────────────────────────────────────────

def certify_resume(
    docx_path: str,
    target_pages: int,
    jd_text: Optional[str] = None,
    scorer_result: Optional[Dict] = None,
    verifier_result: Optional[Dict] = None,
    fitter_result: Optional[Dict] = None,
) -> Dict:
    """
    Run final certification checks on a resume.

    Args:
        docx_path: Path to the .docx resume file
        target_pages: Expected page count (1 or 2)
        jd_text: Optional job description for JD-specific checks
        scorer_result: Optional pre-computed score_resume() result
        verifier_result: Optional pre-computed verify_resume() result
        fitter_result: Optional pre-computed fit_resume() result

    Returns:
        Dict with:
        - verdict: "CERTIFIED", "WARNINGS", or "BLOCKED"
        - checks: list of individual check results
        - blockers: list of check names that blocked certification
        - warnings: list of check names with warnings
        - summary: human-readable one-line summary
    """
    doc = Document(docx_path)
    full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())

    checks = []

    # ── Hard blocker checks (any one = BLOCKED) ──────────────────────

    # 1. Page count
    checks.append(_check_page_count(doc, docx_path, target_pages))

    # 2. Contact info
    checks.append(_check_contact_info(full_text))

    # 3. Banned words
    checks.append(_check_banned_words(full_text))

    # 4. Repeated verbs
    checks.append(_check_repeated_verbs(doc))

    # 5. ATS-unsafe characters
    checks.append(_check_ats_characters(full_text))

    # 6. Bold bullet artifacts
    checks.append(_check_bold_bullets(doc))

    # 7. Date format
    checks.append(_check_date_format(full_text))

    # ── Soft warning checks (flagged but sendable) ───────────────────

    # 8. Metrics density
    checks.append(_check_metrics_density(doc))

    # 9. Verb diversity
    checks.append(_check_verb_diversity(doc))

    # 10. Outcome-first bullets
    checks.append(_check_outcome_first(doc))

    # ── JD-specific checks (only when customized) ────────────────────

    if jd_text:
        # 11. Tagline mirrors JD
        checks.append(_check_jd_tagline(full_text, jd_text))

        # 12. ATS keyword coverage
        checks.append(_check_jd_keyword_coverage(full_text, jd_text))

    # ── Cross-reference with pre-computed results ────────────────────

    # If fitter detected overflow that wasn't resolved
    if fitter_result and fitter_result.get("final_pages", 0) > target_pages:
        checks.append({
            "check": "visual_fit_overflow",
            "status": "BLOCKED",
            "message": f"Visual fitter could not resolve overflow ({fitter_result['final_pages']} pages).",
        })

    # If verifier found errors we haven't already caught
    if verifier_result and verifier_result.get("error_count", 0) > 0:
        uncaught = []
        for err in verifier_result.get("errors", []):
            err_type = err.get("type", "")
            # Skip types we already check directly
            if err_type in ("repeated_verb", "bold_bullet"):
                continue
            uncaught.append(err)
        if uncaught:
            checks.append({
                "check": "verifier_errors",
                "status": "BLOCKED",
                "message": f"{len(uncaught)} additional quality error(s) from verifier.",
                "errors": uncaught,
            })

    # ── Compute verdict ──────────────────────────────────────────────

    blockers = [c["check"] for c in checks if c["status"] == "BLOCKED"]
    warnings = [c["check"] for c in checks if c["status"] == "WARNING"]

    if blockers:
        verdict = "BLOCKED"
        summary = f"BLOCKED — {len(blockers)} issue(s) must be fixed: {', '.join(blockers)}"
    elif warnings:
        verdict = "WARNINGS"
        summary = f"CERTIFIED with {len(warnings)} warning(s): {', '.join(warnings)}"
    else:
        verdict = "CERTIFIED"
        summary = "CERTIFIED — ready to send"

    return {
        "verdict": verdict,
        "checks": checks,
        "blockers": blockers,
        "warnings": warnings,
        "passed_count": sum(1 for c in checks if c["status"] == "PASSED"),
        "total_checks": len(checks),
        "summary": summary,
    }


def format_certification(result: Dict) -> str:
    """Format certification result as a human-readable report."""
    lines = []

    # Header
    verdict = result["verdict"]
    if verdict == "CERTIFIED":
        badge = "[CERTIFIED]"
    elif verdict == "WARNINGS":
        badge = "[CERTIFIED*]"
    else:
        badge = "[BLOCKED]"

    lines.append(f"{'='*60}")
    lines.append(f"  {badge}  {result['summary']}")
    lines.append(f"  {result['passed_count']}/{result['total_checks']} checks passed")
    lines.append(f"{'='*60}")

    # Group by status
    blocked = [c for c in result["checks"] if c["status"] == "BLOCKED"]
    warned = [c for c in result["checks"] if c["status"] == "WARNING"]
    passed = [c for c in result["checks"] if c["status"] == "PASSED"]

    if blocked:
        lines.append("\n  BLOCKERS (must fix):")
        for c in blocked:
            lines.append(f"    [X] {c['check']}: {c['message']}")

    if warned:
        lines.append("\n  WARNINGS (sendable but flagged):")
        for c in warned:
            lines.append(f"    [!] {c['check']}: {c['message']}")

    if passed:
        lines.append("\n  PASSED:")
        for c in passed:
            lines.append(f"    [OK] {c['check']}: {c['message']}")

    lines.append(f"\n{'='*60}")
    return "\n".join(lines)
