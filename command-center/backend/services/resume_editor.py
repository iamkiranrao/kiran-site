"""
Resume Editor Service — XML-level Word document editing with formatting preservation.

Ported from editing_patterns.md and verify_resume.py. Handles:
- Reading template structure
- Rewriting bullets, summaries, skills
- <w:br/> multi-line paragraph management
- Bold/non-bold run management
- Verb uniqueness enforcement
- Quality verification (repeated verbs, bold bullets, metric spacing, bullet length)
"""

import os
import re
import shutil
from typing import Optional, List, Set, Dict, Tuple
from docx import Document
from docx.oxml.ns import qn
from lxml import etree


# ──────────────────────────────────────────────────────────────────────────────
# Core Helper Functions
# ──────────────────────────────────────────────────────────────────────────────

def pt(p) -> str:
    """Get full text of a paragraph, representing <w:br/> as newlines."""
    parts = []
    for child in p._element.iter():
        if child.tag == qn("w:t"):
            parts.append(child.text or "")
        elif child.tag == qn("w:br"):
            parts.append("\n")
    return "".join(parts)


def rewrite(p, new_text: str):
    """Rewrite paragraph text using the first run's formatting.
    For simple single-line paragraphs only."""
    if not p.runs:
        return
    p.runs[0].text = new_text
    for r in p.runs[1:]:
        r.text = ""


def swap(p, old: str, new: str) -> bool:
    """Replace text within a paragraph. Tries run-level first, falls back to full rewrite."""
    for r in p.runs:
        if old in r.text:
            r.text = r.text.replace(old, new)
            return True
    full = pt(p)
    if old in full:
        rewrite(p, full.replace(old, new))
        return True
    return False


def has_br(p) -> bool:
    """Check if a paragraph contains <w:br/> line break elements."""
    return len(p._element.findall(qn("w:r") + "/" + qn("w:br"))) > 0


def get_lines(p) -> List[str]:
    """Split a multi-line paragraph (with <w:br/>) into individual lines."""
    return pt(p).split("\n")


# ──────────────────────────────────────────────────────────────────────────────
# Multi-line Paragraph Editing
# ──────────────────────────────────────────────────────────────────────────────

def rewrite_br_paragraph(p, new_lines: List[str]):
    """Rewrite a paragraph that uses <w:br/> line breaks.
    Preserves formatting from the first text run."""
    if not p.runs:
        return

    # Capture formatting from first text run
    first_run = p.runs[0]
    rpr = first_run._element.find(qn("w:rPr"))
    rpr_xml = etree.tostring(rpr) if rpr is not None else None

    # Remove all existing runs
    for r in p.runs:
        r._element.getparent().remove(r._element)

    # Rebuild with new lines
    p_elem = p._element
    for i, line in enumerate(new_lines):
        if i > 0:
            # Add a <w:br/> before each line after the first
            br_run = etree.SubElement(p_elem, qn("w:r"))
            if rpr_xml:
                br_run.append(etree.fromstring(rpr_xml))
            etree.SubElement(br_run, qn("w:br"))

        # Add the text run
        new_run = etree.SubElement(p_elem, qn("w:r"))
        if rpr_xml:
            new_run.append(etree.fromstring(rpr_xml))
        t_elem = etree.SubElement(new_run, qn("w:t"))
        t_elem.text = line
        if line.startswith(" ") or line.endswith(" "):
            t_elem.set(qn("xml:space"), "preserve")


# ──────────────────────────────────────────────────────────────────────────────
# Bold Management
# ──────────────────────────────────────────────────────────────────────────────

def ensure_not_bold(run):
    """Explicitly set a run to non-bold."""
    rpr = run._element.find(qn("w:rPr"))
    if rpr is None:
        rpr = etree.SubElement(run._element, qn("w:rPr"))
        run._element.insert(0, rpr)

    for b in rpr.findall(qn("w:b")):
        rpr.remove(b)

    b_elem = etree.SubElement(rpr, qn("w:b"))
    b_elem.set(qn("w:val"), "0")


def ensure_bold(run):
    """Explicitly set a run to bold."""
    rpr = run._element.find(qn("w:rPr"))
    if rpr is None:
        rpr = etree.SubElement(run._element, qn("w:rPr"))
        run._element.insert(0, rpr)

    for b in rpr.findall(qn("w:b")):
        rpr.remove(b)

    etree.SubElement(rpr, qn("w:b"))


def bold_sweep(doc):
    """Post-edit sweep: ensure all bullet text runs in <w:br/> paragraphs are NOT bold.
    Only affects runs that contain actual text (not headers)."""
    for p in doc.paragraphs:
        if not has_br(p):
            continue

        lines = get_lines(p)
        # Skip the header paragraph (paragraph 0 with name/contact)
        if any(line.strip().isupper() and len(line.strip()) < 40 for line in lines):
            continue

        # Find runs that contain bullet text (not headers, not <w:br/> runs)
        for run in p.runs:
            text = run.text.strip() if run.text else ""
            if not text:
                continue
            # If text is > 30 chars or starts with bullet marker, it's bullet text
            if len(text) > 30 or text.startswith("•") or text.startswith("-"):
                if run.bold:
                    ensure_not_bold(run)


# ──────────────────────────────────────────────────────────────────────────────
# Section Finders
# ──────────────────────────────────────────────────────────────────────────────

def find_tagline(doc) -> Tuple[Optional[int], Optional[object]]:
    """Find the italic tagline paragraph (appears after contact info, before first section).

    The tagline is identified by:
    - Appears in the first 10 paragraphs
    - All runs are italic
    - Length between 30 and 200 characters
    - Not a section header (not all-bold, not all-caps)
    - Not the summary paragraph (which comes after a SUMMARY header)
    """
    for i, p in enumerate(doc.paragraphs):
        if i > 10:
            break
        text = pt(p).strip()
        if not text or len(text) < 30 or len(text) > 200:
            continue
        # Skip section headers (all-caps or all-bold)
        if text.upper() == text and len(text) < 40:
            continue
        if not p.runs:
            continue
        # Check if all non-empty runs are italic
        non_empty_runs = [r for r in p.runs if r.text and r.text.strip()]
        if not non_empty_runs:
            continue
        if all(r.italic for r in non_empty_runs):
            return i, p
    return None, None


def find_summary(doc) -> Tuple[Optional[int], Optional[object]]:
    """Find the professional summary paragraph."""
    # First, look for a SUMMARY section header and return the paragraph after it
    for i, p in enumerate(doc.paragraphs):
        text = pt(p).strip()
        if text.upper().replace(" ", "") in ("SUMMARY", "PROFESSIONALSUMMARY", "EXECUTIVESUMMARY"):
            # Return the next non-empty paragraph
            for j in range(i + 1, min(i + 4, len(doc.paragraphs))):
                next_text = pt(doc.paragraphs[j]).strip()
                if next_text and len(next_text) > 30:
                    return j, doc.paragraphs[j]
            break

    # Fallback: find a long, non-bold, non-BR paragraph in the first 15 paragraphs.
    # Only use this fallback if there's a reasonable chance of a summary paragraph
    # (i.e., a long, plain-text, non-header, non-bullet paragraph).
    # Skip paragraphs that look like contact info (contain phone/email markers).
    contact_markers = ["@", "📞", "☎", "✉", "linkedin", "portfolio", "github"]
    for i, p in enumerate(doc.paragraphs):
        text = pt(p).strip()
        if i > 15:
            break
        if len(text) < 50 or has_br(p):
            continue
        # Skip bullet lines (skills, competencies, etc.)
        if text.startswith("•") or text.startswith("-"):
            continue
        # Skip contact info lines
        text_lower = text.lower()
        if any(marker in text_lower for marker in contact_markers):
            continue
        # Skip any paragraph with a bold first run (role headers, section headers)
        if p.runs and p.runs[0].bold:
            continue
        # Skip all-italic paragraphs (taglines)
        non_empty_runs = [r for r in p.runs if r.text and r.text.strip()]
        if non_empty_runs and all(r.italic for r in non_empty_runs):
            continue
        # Skip paragraphs containing tab characters (role header format: "Title | Company\tDate")
        if "\t" in text:
            continue
        # This is likely the summary
        return i, p

    return None, None


def find_section_header(doc, section_name: str) -> Optional[int]:
    """Find the paragraph index of a section header (e.g., 'EXPERIENCE', 'SKILLS')."""
    for i, p in enumerate(doc.paragraphs):
        text = pt(p).strip()
        if text.upper() == section_name.upper():
            return i
    return None


def find_job_section(doc, company_name: str) -> Tuple[Optional[int], Optional[int]]:
    """Find the paragraph index range for a specific job entry."""
    start = None
    for i, p in enumerate(doc.paragraphs):
        text = pt(p)
        if company_name.lower() in text.lower():
            start = i
            break

    if start is None:
        return None, None

    end = len(doc.paragraphs)
    for i in range(start + 1, len(doc.paragraphs)):
        text = pt(doc.paragraphs[i]).strip()
        if text and len(text) < 40:
            all_bold = all(r.bold for r in doc.paragraphs[i].runs if r.text.strip())
            if all_bold:
                end = i
                break

    return start, end


def find_bullet_paragraphs(doc, company_name: str) -> Tuple[Optional[int], List[int]]:
    """Find the company header and its individual bullet paragraphs.

    Handles two template styles:
    1. BR-style: all bullets in one paragraph separated by <w:br/>
    2. Individual-style: each bullet is its own paragraph (Detailed templates)

    Returns (header_index, [bullet_paragraph_indices]).
    For BR-style, returns a single-element list with the BR paragraph index.
    """
    # Find the company header paragraph
    header_idx = None
    for i, p in enumerate(doc.paragraphs):
        text = pt(p)
        if company_name.lower() in text.lower():
            header_idx = i
            break

    if header_idx is None:
        return None, []

    # Check if the paragraph right after the header (or within 2) is a BR-style bullet block
    for j in range(header_idx + 1, min(header_idx + 3, len(doc.paragraphs))):
        p = doc.paragraphs[j]
        text = pt(p).strip()
        if not text:
            continue
        if has_br(p) and len(text) > 50:
            # Check it's NOT another company header (bold + short per line)
            all_bold = all(r.bold for r in p.runs if r.text and r.text.strip()) if p.runs else False
            if not all_bold:
                return header_idx, [j]  # BR-style: single paragraph with all bullets

    # Individual-style: collect non-bold, non-empty paragraphs after the header
    bullet_indices = []
    for j in range(header_idx + 1, len(doc.paragraphs)):
        p = doc.paragraphs[j]
        text = pt(p).strip()

        if not text:
            # Empty paragraph — stop if we already found bullets
            if bullet_indices:
                break
            continue

        # Check if this is a new section/company header (all bold)
        all_bold = all(r.bold for r in p.runs if r.text and r.text.strip()) if p.runs else False
        if all_bold:
            break

        # Check if this is another company entry with BR (e.g., "[BR]VP OF PRODUCT...")
        if has_br(p) and any(r.bold for r in p.runs if r.text and r.text.strip()):
            break

        # Check if this is a new role header (mixed bold: bold title + separator + date)
        # Handles clean templates where only the job title is bold, not the full line
        has_bold = any(r.bold for r in p.runs if r.text and r.text.strip()) if p.runs else False
        if has_bold and not all_bold and ("|" in text or "\u2013" in text or "\u2014" in text):
            if re.search(r'(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+\d{4}', text):
                break

        bullet_indices.append(j)

    return header_idx, bullet_indices


# ──────────────────────────────────────────────────────────────────────────────
# Verb Tracking
# ──────────────────────────────────────────────────────────────────────────────

VERB_BANK = [
    # Curated per RESUME-RULES.md — excludes hard-banned verbs (Orchestrated,
    # Spearheaded, Championed, Pioneered) and soft-banned verbs (Architected,
    # Streamlined, Optimized, Transformed) to minimize AI-detection flags.
    "Accelerated", "Achieved", "Activated", "Advanced", "Aligned",
    "Amplified", "Automated", "Boosted", "Built",
    "Catalyzed", "Co-created", "Consolidated", "Converted",
    "Coordinated", "Crafted", "Cultivated", "Customized", "Decreased",
    "Defined", "Delivered", "Designed", "Developed", "Directed",
    "Drove", "Elevated", "Eliminated", "Enabled", "Engineered",
    "Established", "Evaluated", "Evolved", "Executed", "Expanded",
    "Facilitated", "Forged", "Founded", "Generated", "Grew",
    "Guided", "Halved", "Headed", "Identified", "Implemented",
    "Improved", "Increased", "Influenced", "Initiated", "Innovated",
    "Integrated", "Introduced", "Launched", "Led", "Managed",
    "Mapped", "Mentored", "Migrated", "Modernized", "Negotiated",
    "Operationalized", "Overhauled", "Partnered",
    "Piloted", "Prioritized", "Produced", "Propelled",
    "Rationalized", "Rebuilt", "Redesigned", "Reduced", "Refined",
    "Reimagined", "Replatformed", "Restructured", "Revamped", "Reversed",
    "Scaled", "Secured", "Shaped", "Simplified",
    "Standardized", "Steered", "Strengthened", "Surpassed",
    "Systematized", "Tripled", "Unified", "Unlocked",
]


class VerbTracker:
    """Track used action verbs across the entire resume to prevent duplicates."""

    def __init__(self):
        self.used: Set[str] = set()

    def pick_verb(self, preferred: str, alternatives: Optional[List[str]] = None) -> str:
        """Return preferred verb if unused, otherwise try alternatives, then verb bank."""
        if preferred.lower() not in self.used:
            self.used.add(preferred.lower())
            return preferred

        if alternatives:
            for alt in alternatives:
                if alt.lower() not in self.used:
                    self.used.add(alt.lower())
                    return alt

        # Fall back to verb bank
        for verb in VERB_BANK:
            if verb.lower() not in self.used:
                self.used.add(verb.lower())
                return verb

        return preferred  # Last resort

    def register(self, verb: str):
        """Register a verb as used without picking."""
        self.used.add(verb.lower())


def get_experience_bullet_map(doc) -> List[Dict]:
    """Extract company headers and their bullet counts from the template.

    Returns a list like:
    [
        {"header_text": "VP OF PRODUCT | Wells Fargo Strategy...", "bullet_count": 6, "header_idx": 23},
        ...
    ]
    Used to tell Claude exactly how many bullets to generate per company.
    """
    result = []

    # Find the EXPERIENCE section — may be a standalone header or combined with first company
    exp_start = None
    for i, p in enumerate(doc.paragraphs):
        text = pt(p).strip()
        text_upper = text.upper()
        # Match "EXPERIENCE" or "PROFESSIONAL EXPERIENCE" (standalone or combined with first company)
        if ("EXPERIENCE" in text_upper and
                any(text_upper.startswith(prefix) for prefix in
                    ("EXPERIENCE", "PROFESSIONAL EXPERIENCE", "WORK EXPERIENCE"))):
            exp_start = i
            break

    if exp_start is None:
        return result

    # Find all company headers: bold paragraphs within the experience section
    # Stop at the next major section (EDUCATION, SKILLS, etc.)
    stop_sections = {"EDUCATION", "EDUCATION & CERTIFICATIONS", "TECHNICAL SKILLS",
                     "SKILLS", "CORE SKILLS", "CERTIFICATIONS", "TOOLS",
                     "LEADERSHIP & STRATEGIC COMPETENCIES",
                     "CORE COMPETENCIES", "TECHNICAL COMPETENCIES"}

    company_headers = []
    for i in range(exp_start, len(doc.paragraphs)):
        p = doc.paragraphs[i]
        text = pt(p).strip()
        if not text:
            continue

        # Check if we've hit a non-experience section
        text_upper = text.upper().strip()
        if any(text_upper == s or text_upper.startswith(s + "\n") for s in stop_sections):
            if i > exp_start:
                break

        # Detect company headers: bold, contains separators and date-like patterns
        has_bold = any(r.bold for r in p.runs if r.text and r.text.strip()) if p.runs else False
        has_separator = "|" in text or "–" in text or "—" in text
        has_date = bool(re.search(r'(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+\d{4}', text) or
                        re.search(r"('{0,1}\d{2,4})\s*[–—-]\s*(Present|\d{2,4})", text))
        if has_bold and has_separator and has_date and len(text) > 20:
            company_headers.append(i)

    # For each company header, use a distinctive substring to find bullets
    for header_i in company_headers:
        header_text = pt(doc.paragraphs[header_i]).replace("\n", " ").strip()
        # Clean up extra whitespace from the header
        header_text_clean = re.sub(r'\s{2,}', ' ', header_text)

        # Try each pipe-separated segment to find the one that locates bullets
        parts = [p.strip() for p in header_text.split("|")]
        bullet_indices = []

        for part in parts:
            part_clean = part.split(",")[0].strip()
            if len(part_clean) < 3:
                continue
            _, found = find_bullet_paragraphs(doc, part_clean)
            if found and len(found) > len(bullet_indices):
                bullet_indices = found

        result.append({
            "header_idx": header_i,
            "header_text": header_text_clean[:120],
            "bullet_count": len(bullet_indices),
        })

    return result


# ──────────────────────────────────────────────────────────────────────────────
# Skills Editing
# ──────────────────────────────────────────────────────────────────────────────

def reorder_skills(p, priority_skills: List[str]):
    """Move JD-relevant skills to the front of a delimited skills line.
    Supports both comma-separated and pipe-separated (│) formats."""
    text = pt(p)

    # Detect delimiter: pipe (│) or comma
    if "│" in text:
        delimiter = "│"
    elif "|" in text:
        delimiter = "|"
    else:
        delimiter = ","

    skills = [s.strip() for s in text.split(delimiter)]

    priority = [s for s in skills if any(ps.lower() in s.lower() for ps in priority_skills)]
    other = [s for s in skills if s not in priority]

    joiner = f" {delimiter} " if delimiter in ("│", "|") else ", "
    new_text = joiner.join(priority + other)
    rewrite(p, new_text)


def write_core_competencies(doc, skills: List[str]) -> bool:
    """Replace the Core Competencies section with new skills from the skills library.

    Core Competencies layout: 4 paragraphs, each with 4 bullet-formatted skills.
    Format per row: '• Skill 1  • Skill 2  • Skill 3  • Skill 4'

    Returns True if written, False if section not found.
    """
    # Find the Core Competencies header
    header_idx = None
    for i, p in enumerate(doc.paragraphs):
        text = pt(p).strip()
        if text.upper() in ("CORE COMPETENCIES", "LEADERSHIP & STRATEGIC COMPETENCIES"):
            header_idx = i
            break

    if header_idx is None:
        return False

    # Find the competency rows (next 4 paragraphs with bullet content)
    row_indices = []
    for j in range(header_idx + 1, min(header_idx + 8, len(doc.paragraphs))):
        text = pt(doc.paragraphs[j]).strip()
        if not text:
            continue
        # Stop at section headers
        text_upper = text.upper()
        if text_upper in ("EXPERIENCE", "PROFESSIONAL EXPERIENCE", "WORK EXPERIENCE",
                          "EDUCATION", "TECHNICAL COMPETENCIES", "TECHNICAL SKILLS"):
            break
        if "•" in text:
            row_indices.append(j)

    if not row_indices:
        return False

    # Build 4 rows of 4 skills each
    padded = skills[:16]
    while len(padded) < 16:
        padded.append("")

    rows = []
    for r in range(4):
        row_skills = padded[r * 4 : (r + 1) * 4]
        row_text = "  ".join(f"• {s}" for s in row_skills if s)
        rows.append(row_text)

    # Write to existing paragraphs (reuse formatting)
    for k, row_idx in enumerate(row_indices):
        if k < len(rows):
            rewrite(doc.paragraphs[row_idx], rows[k])

    return True


def write_technical_skills(doc, tech_data, version: str) -> bool:
    """Replace the Technical Competencies/Skills section with new skills from the skills library.

    For 1-Page: tech_data is a flat list → single pipe-delimited paragraph
    For 2-Page/Detailed: tech_data is a dict of category→[skills] → one paragraph per category

    Returns True if written, False if section not found.
    """
    # Find the Technical Competencies/Skills header
    header_idx = None
    for i, p in enumerate(doc.paragraphs):
        text = pt(p).strip()
        if text.upper() in ("TECHNICAL COMPETENCIES", "TECHNICAL SKILLS", "TOOLS"):
            header_idx = i
            break

    if header_idx is None:
        return False

    # Find the skills paragraphs after the header
    skill_indices = []
    for j in range(header_idx + 1, min(header_idx + 10, len(doc.paragraphs))):
        p = doc.paragraphs[j]
        text = pt(p).strip()
        if not text:
            break
        # Stop at section headers (bold + short + no delimiters)
        all_bold = all(r.bold for r in p.runs if r.text and r.text.strip()) if p.runs else False
        if all_bold and len(text) < 40 and "|" not in text and "," not in text:
            break
        if "|" in text or "," in text or ":" in text:
            skill_indices.append(j)

    if not skill_indices:
        return False

    if version == "1-Page":
        # Flat format: single pipe-delimited line
        if isinstance(tech_data, list):
            new_text = " | ".join(tech_data)
        else:
            # If dict passed for 1-page, flatten
            all_skills = []
            for cat_skills in tech_data.values():
                all_skills.extend(cat_skills)
            new_text = " | ".join(all_skills[:22])

        # Write to first skills paragraph, clear the rest
        rewrite(doc.paragraphs[skill_indices[0]], new_text)
        for idx in skill_indices[1:]:
            rewrite(doc.paragraphs[idx], "")

    else:
        # Grouped format: 'Category: Skill | Skill | ...'
        if isinstance(tech_data, dict):
            lines = []
            for cat, skills in tech_data.items():
                lines.append(f"{cat}: {' | '.join(skills)}")
        else:
            # Flat list passed for 2-page, write as single line
            lines = [" | ".join(tech_data)]

        # Write lines to existing paragraphs
        for k, idx in enumerate(skill_indices):
            if k < len(lines):
                rewrite(doc.paragraphs[idx], lines[k])
            else:
                rewrite(doc.paragraphs[idx], "")

    return True


# ──────────────────────────────────────────────────────────────────────────────
# Template Management
# ──────────────────────────────────────────────────────────────────────────────

TEMPLATES_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "templates")

TEMPLATE_MAP = {
    ("PM", "1-Page"): "PM_1Pager.docx",
    ("PM", "2-Page"): "PM_2Pager.docx",
    ("PM", "Detailed"): "PM_Detailed.docx",
    ("PjM", "1-Page"): "PjM_1Pager.docx",
    ("PjM", "2-Page"): "PjM_2Pager.docx",
    ("PjM", "Detailed"): "PjM_Detailed.docx",
    ("PMM", "1-Page"): "PMM_1Pager.docx",
    ("PMM", "2-Page"): "PMM_2Pager.docx",
    ("PMM", "Detailed"): "PMM_Detailed.docx",
}


def detect_persona(jd_text: str) -> str:
    """Auto-detect persona from job description text."""
    jd_lower = jd_text.lower()
    pm_signals = ["product manager", "group pm", "apm", "head of product", "director of product"]
    pjm_signals = ["program manager", "project manager", "tpm", "delivery manager", "technical program"]
    pmm_signals = ["product marketing", "pmm", "growth marketing", "gtm", "go-to-market"]

    pm_score = sum(1 for s in pm_signals if s in jd_lower)
    pjm_score = sum(1 for s in pjm_signals if s in jd_lower)
    pmm_score = sum(1 for s in pmm_signals if s in jd_lower)

    if pmm_score > pm_score and pmm_score > pjm_score:
        return "PMM"
    if pjm_score > pm_score:
        return "PjM"
    return "PM"


def copy_template(persona: str, version: str, dest_path: str) -> str:
    """Copy the selected template to a working location. Never modify originals."""
    key = (persona, version)
    if key not in TEMPLATE_MAP:
        raise ValueError(f"No template for persona={persona}, version={version}")

    src = os.path.join(TEMPLATES_DIR, TEMPLATE_MAP[key])
    if not os.path.exists(src):
        raise FileNotFoundError(f"Template not found: {src}")

    # Use copy (not copy2) to avoid inheriting restrictive source permissions,
    # then explicitly ensure the working copy is writable
    shutil.copy(src, dest_path)
    os.chmod(dest_path, 0o644)
    return dest_path


def read_template_structure(docx_path: str) -> dict:
    """Read and analyze the template structure for editing guidance."""
    doc = Document(docx_path)
    structure = {
        "total_paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "sections": [],
        "bullet_paragraphs": [],
        "summary_index": None,
    }

    for i, p in enumerate(doc.paragraphs):
        text = pt(p).strip()
        if not text:
            continue

        br_count = len(p._element.findall(".//" + qn("w:br")))
        all_bold = all(r.bold for r in p.runs if r.text.strip()) if p.runs else False

        if all_bold and len(text) < 40:
            structure["sections"].append({"index": i, "text": text})

        if br_count > 0 and len(text) > 50:
            lines = text.split("\n")
            structure["bullet_paragraphs"].append({
                "index": i,
                "line_count": len([l for l in lines if l.strip()]),
                "preview": lines[0][:80] if lines else "",
            })

    # Find summary
    idx, _ = find_summary(doc)
    structure["summary_index"] = idx

    return structure


def extract_template_content(doc) -> dict:
    """Extract the current text from each editable section of a resume template.

    Returns a dict like:
    {
        "sections": [
            {
                "id": "summary",
                "label": "Professional Summary",
                "current_text": "Kiran is a product leader..."
            },
            {
                "id": "career_highlights",
                "label": "Career Highlights",
                "current_text": ["• Led mobile banking...", "• Drove AI integration..."]
            },
            {
                "id": "exp__Wells Fargo Strategy",
                "label": "Wells Fargo Strategy",
                "header_text": "VP OF PRODUCT | Wells Fargo Strategy...",
                "current_text": ["• Led redesign of...", "• Managed $12M..."]
            },
            ...
            {
                "id": "skills",
                "label": "Skills",
                "current_text": "Product Strategy │ Agile │ ..."
            }
        ]
    }
    """
    sections = []

    # ── Tagline ──────────────────────────────────────────────────────────────
    tagline_idx, tagline_p = find_tagline(doc)
    if tagline_p:
        sections.append({
            "id": "tagline",
            "label": "Tagline",
            "current_text": pt(tagline_p).strip(),
        })

    # ── Summary ─────────────────────────────────────────────────────────────
    summary_idx, summary_p = find_summary(doc)
    if summary_p:
        sections.append({
            "id": "summary",
            "label": "Professional Summary",
            "current_text": pt(summary_p).strip(),
        })

    # ── Career Highlights ───────────────────────────────────────────────────
    for i, p in enumerate(doc.paragraphs):
        text = pt(p).strip()
        if text.upper() == "CAREER HIGHLIGHTS":
            # Check for BR-joined paragraph first, then individual bullet paragraphs
            lines = []
            for j in range(i + 1, min(i + 8, len(doc.paragraphs))):
                ch_p = doc.paragraphs[j]
                ch_text = pt(ch_p).strip()
                if not ch_text:
                    break
                # Stop at section headers (bold, short, no bullets)
                ch_all_bold = all(r.bold for r in ch_p.runs if r.text and r.text.strip()) if ch_p.runs else False
                if ch_all_bold and len(ch_text) < 40 and "•" not in ch_text:
                    break
                if has_br(ch_p):
                    # BR-joined paragraph — split into lines
                    lines = [l.strip() for l in ch_text.split("\n") if l.strip()]
                    break
                if ch_text.startswith("•") or ch_text.startswith("-"):
                    lines.append(ch_text)
            if lines:
                sections.append({
                    "id": "career_highlights",
                    "label": "Career Highlights",
                    "current_text": lines,
                })
            break

    # ── Experience — per company ────────────────────────────────────────────
    bullet_map = get_experience_bullet_map(doc)
    for entry in bullet_map:
        header_text = entry["header_text"]
        header_idx = entry["header_idx"]

        # Try to find bullets using company name substrings
        parts = [p.strip() for p in header_text.split("|")]
        bullet_indices = []
        matching_key = ""

        for part in parts:
            part_clean = part.split(",")[0].strip()
            if len(part_clean) < 3:
                continue
            _, found = find_bullet_paragraphs(doc, part_clean)
            if found and len(found) > len(bullet_indices):
                bullet_indices = found
                matching_key = part_clean

        if not bullet_indices:
            continue

        # Extract bullet text
        if len(bullet_indices) == 1 and has_br(doc.paragraphs[bullet_indices[0]]):
            lines = [l.strip() for l in pt(doc.paragraphs[bullet_indices[0]]).split("\n") if l.strip()]
        else:
            lines = [pt(doc.paragraphs[bi]).strip() for bi in bullet_indices]

        # Build a short label from the header
        label = matching_key if matching_key else header_text[:40]

        sections.append({
            "id": f"exp__{matching_key}",
            "label": label,
            "header_text": header_text,
            "current_text": lines,
        })

    # ── Core Competencies ────────────────────────────────────────────────────
    core_headers = {"CORE COMPETENCIES", "LEADERSHIP & STRATEGIC COMPETENCIES", "CORE SKILLS"}
    for i, p in enumerate(doc.paragraphs):
        text = pt(p).strip()
        if text.upper() in core_headers:
            core_lines = []
            for j in range(i + 1, min(i + 8, len(doc.paragraphs))):
                core_p = doc.paragraphs[j]
                core_text = pt(core_p).strip()
                if not core_text:
                    break
                # Stop at section headers (bold, short text)
                c_all_bold = all(r.bold for r in core_p.runs if r.text and r.text.strip()) if core_p.runs else False
                if c_all_bold and len(core_text) < 40:
                    break
                # Core competencies use bullet (•) separators
                if "•" in core_text:
                    core_lines.append(core_text)
            if core_lines:
                sections.append({
                    "id": "core_competencies",
                    "label": text,
                    "current_text": core_lines,
                })
            break

    # ── Technical Competencies ────────────────────────────────────────────────
    tech_headers = {"TECHNICAL COMPETENCIES", "TECHNICAL SKILLS", "TOOLS", "SKILLS"}
    for i, p in enumerate(doc.paragraphs):
        text = pt(p).strip()
        if text.upper() in tech_headers:
            tech_lines = []
            for j in range(i + 1, min(i + 8, len(doc.paragraphs))):
                tech_p = doc.paragraphs[j]
                tech_text = pt(tech_p).strip()
                if not tech_text:
                    break
                # Stop at section headers (bold, short text)
                t_all_bold = all(r.bold for r in tech_p.runs if r.text and r.text.strip()) if tech_p.runs else False
                if t_all_bold and len(tech_text) < 40:
                    break
                # Technical skills use pipe (|) or colon (:) separators
                if "|" in tech_text or "│" in tech_text or ":" in tech_text or "," in tech_text:
                    tech_lines.append(tech_text)
            if tech_lines:
                sections.append({
                    "id": "technical_competencies",
                    "label": text,
                    "current_text": tech_lines if len(tech_lines) > 1 else tech_lines[0],
                })
            break

    return {"sections": sections}


# ──────────────────────────────────────────────────────────────────────────────
# Quality Verification
# ──────────────────────────────────────────────────────────────────────────────

def verify_resume(docx_path: str, check_interests: bool = False) -> dict:
    """Run all quality checks. Returns dict with errors, warnings, and details."""
    doc = Document(docx_path)
    errors = []
    warnings = []

    # 1. Repeated verbs
    verb_locations: Dict[str, List] = {}
    for i, p in enumerate(doc.paragraphs):
        text = pt(p)
        lines = text.split("\n")
        for line in lines:
            line = line.strip()
            if not line or len(line) < 15:
                continue
            words = line.split()
            if not words:
                continue
            verb = words[0].rstrip(",.:;")
            if verb[0].isdigit() or len(verb) < 3:
                continue
            if verb.isupper() and len(verb) > 3:
                continue
            # Skip education keywords
            if verb.lower() in ("bachelor", "master", "mba", "certified", "university", "college"):
                continue
            verb_lower = verb.lower()
            if verb_lower not in verb_locations:
                verb_locations[verb_lower] = []
            verb_locations[verb_lower].append({"paragraph": i, "text": line[:80]})

    for verb, locs in verb_locations.items():
        if len(locs) > 1:
            errors.append({
                "type": "repeated_verb",
                "verb": verb,
                "count": len(locs),
                "locations": locs,
            })

    # 2. Bold bullets
    for i, p in enumerate(doc.paragraphs):
        br_count = len(p._element.findall(".//" + qn("w:br")))
        if br_count == 0:
            continue
        for j, run in enumerate(p.runs):
            if run.bold and run.text and len(run.text.strip()) > 30:
                errors.append({
                    "type": "bold_bullet",
                    "paragraph": i,
                    "run": j,
                    "text": run.text[:60],
                })

    # 3. Metric spacing
    bad_pattern = re.compile(r"\d\s+[MKB%](?:\b|$)")
    for i, p in enumerate(doc.paragraphs):
        text = pt(p)
        matches = bad_pattern.findall(text)
        if matches:
            warnings.append({
                "type": "metric_spacing",
                "paragraph": i,
                "matches": matches,
            })

    # 4. Bullet length
    for i, p in enumerate(doc.paragraphs):
        text = pt(p)
        for line_num, line in enumerate(text.split("\n")):
            line = line.strip()
            if len(line) > 120:
                warnings.append({
                    "type": "long_bullet",
                    "paragraph": i,
                    "line": line_num,
                    "length": len(line),
                    "text": line[:80] + "...",
                })

    # 5. Interests section (removed from templates — no longer checked)

    return {
        "errors": errors,
        "warnings": warnings,
        "error_count": len(errors),
        "warning_count": len(warnings),
        "passed": len(errors) == 0,
    }
