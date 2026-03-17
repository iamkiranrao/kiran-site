"""
Resume Fitter — Post-customization visual fitting to eliminate whitespace and prevent overflow.


Runs after content edits and quality checks. Measures actual page fill via PDF conversion,
then applies visual adjustments in a strict priority order to achieve optimal fill.

Two modes:
1. FILL mode (page fill < 85%): expand spacing/margins/font to fill the page
2. TRIM mode (page overflow): tighten spacing/margins/font to fit within page limit

Adjustment hierarchy (applied one at a time, re-measured after each):

FILL mode (too much whitespace):
  1. Increase section header spacing (space_before on section headers)
  2. Increase paragraph spacing (space_after on bullet paragraphs)
  3. Increase margins (reduces text area, content fills proportionally more)
  4. Increase body font size (last resort — subtle, 0.5pt max)

TRIM mode (overflow):
  1. Decrease paragraph spacing (space_after on bullet paragraphs)
  2. Decrease section header spacing (space_before on section headers)
  3. Decrease margins (expands text area, fits more per line)
  4. Decrease body font size (last resort — 9.5pt minimum)

Section-split prevention:
  - If a section header lands alone at the bottom of a page, add keepWithNext
    to ensure it moves to the next page with its first bullet.

Hard bounds:
  - Margins: 0.4" to 0.8" (all four sides move together)
  - Body font: 9.5pt to 10.5pt
  - Section spacing: 2pt to 14pt (space_before on headers)
  - Paragraph spacing: 0pt to 4pt (space_after on bullets)
  - Line spacing: not modified (too disruptive to the design system)

Usage:
    from services.resume_fitter import fit_resume
    result = fit_resume("/path/to/resume.docx", target_pages=2)
"""

import os
import re
import subprocess
import logging
import tempfile
from typing import Optional, Dict, List, Tuple
from docx import Document
from docx.shared import Inches, Pt, Emu, Twips
from docx.oxml.ns import qn
from lxml import etree


# ── Constants ──────────────────────────────────────────────────────────

# Page fill targets
FILL_LOW = 0.85   # Below this = too much whitespace, enter FILL mode
FILL_HIGH = 1.00  # At or above = overflow, enter TRIM mode
FILL_IDEAL = 0.92 # Target fill percentage

# Margin bounds (inches) — all four sides move together
MARGIN_MIN = 0.40
MARGIN_MAX = 0.80
MARGIN_STEP = 0.05  # Adjust by 0.05" per iteration

# Body font bounds (points)
FONT_MIN = Pt(9.5)
FONT_MAX = Pt(10.5)
FONT_STEP = Pt(0.5)  # Adjust by 0.5pt per iteration

# Section header spacing bounds (points, space_before)
SECTION_SPACING_MIN = Pt(2)
SECTION_SPACING_MAX = Pt(14)
SECTION_SPACING_STEP = Pt(1.5)  # Adjust by 1.5pt per iteration

# Paragraph spacing bounds (points, space_after on bullets)
PARA_SPACING_MIN = Pt(0)
PARA_SPACING_MAX = Pt(4)
PARA_SPACING_STEP = Pt(0.5)  # Adjust by 0.5pt per iteration

# Max adjustment iterations to prevent infinite loops
MAX_ITERATIONS = 8

# Section headers that should never be orphaned at page bottom
SECTION_HEADERS = {
    "EXPERIENCE", "PROFESSIONAL EXPERIENCE", "EDUCATION",
    "EDUCATION & CERTIFICATIONS", "SKILLS", "CORE SKILLS",
    "TECHNICAL SKILLS", "CORE COMPETENCIES", "LEADERSHIP & STRATEGIC COMPETENCIES",
    "CAREER HIGHLIGHTS", "INTERESTS", "CERTIFICATIONS",
}


# ── Page Measurement ──────────────────────────────────────────────────

def _convert_to_pdf(docx_path: str) -> Optional[str]:
    """Convert docx to PDF using LibreOffice headless. Returns PDF path or None."""
    outdir = os.path.dirname(docx_path) or tempfile.gettempdir()
    try:
        result = subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", outdir, docx_path],
            capture_output=True, timeout=30,
        )
        pdf_path = docx_path.rsplit(".", 1)[0] + ".pdf"
        if os.path.exists(pdf_path):
            return pdf_path
    except Exception:
        logger.warning("LibreOffice PDF conversion failed", exc_info=True)
    return None


def _count_pdf_pages(pdf_path: str) -> int:
    """Count pages in a PDF file."""
    try:
        import PyPDF2
        reader = PyPDF2.PdfReader(pdf_path)
        return len(reader.pages)
    except Exception:
        logger.warning("PDF page count failed", exc_info=True)
        return 0


def _estimate_page_fill(docx_path: str, target_pages: int) -> Tuple[float, int]:
    """
    Estimate how full the last page is by converting to PDF.

    Returns (fill_ratio, actual_pages) where:
    - fill_ratio: 0.0-1.0+ indicating how full the content is relative to target
    - actual_pages: number of pages the PDF actually rendered to

    If actual_pages > target_pages, fill_ratio > 1.0 (overflow).
    If actual_pages == target_pages, fill_ratio is estimated from content and visual properties.
    If actual_pages < target_pages, fill_ratio < 1.0 (underfill).
    """
    pdf_path = _convert_to_pdf(docx_path)
    if not pdf_path:
        return 0.9, target_pages  # Default: assume roughly filled, don't adjust

    actual_pages = _count_pdf_pages(pdf_path)

    # Clean up PDF
    try:
        os.remove(pdf_path)
    except Exception:
        logger.debug("Could not clean up temp PDF: %s", pdf_path)

    if actual_pages > target_pages:
        # Overflow — fill > 1.0
        return 1.0 + (actual_pages - target_pages) * 0.1, actual_pages
    elif actual_pages < target_pages:
        # Significant underfill — content dropped below page count
        return actual_pages / target_pages, actual_pages
    else:
        # Right page count — estimate fill from content density + visual properties
        doc = Document(docx_path)
        total_words = sum(len(p.text.split()) for p in doc.paragraphs if p.text.strip())

        # Base word capacity per page
        if target_pages == 1:
            base_capacity = 425
        else:
            base_capacity = target_pages * 375

        # Adjust capacity based on visual properties that reduce usable space
        # Larger margins = less text area = lower effective capacity
        for section in doc.sections:
            margin_inches = section.left_margin.inches
            # Baseline margins: 0.6" (1-page) or 0.7" (2-page)
            baseline_margin = 0.6 if target_pages == 1 else 0.7
            margin_delta = margin_inches - baseline_margin
            # Each 0.05" of extra margin reduces capacity by ~3%
            base_capacity *= (1 - margin_delta * 0.6)
            break  # All sections same in our templates

        # Larger spacing = each paragraph takes more room = lower effective capacity
        spacing_overhead = 0
        for p in doc.paragraphs:
            pf = p.paragraph_format
            if pf.space_before and pf.space_before > Pt(8):
                # Extra spacing above baseline 8pt
                spacing_overhead += (pf.space_before.pt - 8) * 0.3
            if pf.space_after and pf.space_after > Pt(1):
                spacing_overhead += (pf.space_after.pt - 1) * 0.15

        # Convert spacing overhead to word-equivalent reduction
        base_capacity -= spacing_overhead

        # Larger font = fewer words per line = lower capacity
        body_fonts = []
        for p in doc.paragraphs:
            for r in p.runs:
                if r.font.size and 9.0 <= r.font.size.pt <= 11.0 and r.font.size.pt not in (11.0, 20.0):
                    body_fonts.append(r.font.size.pt)
        if body_fonts:
            avg_font = sum(body_fonts) / len(body_fonts)
            # Baseline is 10pt; each 0.5pt increase reduces capacity by ~5%
            font_delta = avg_font - 10.0
            base_capacity *= (1 - font_delta * 0.1)

        base_capacity = max(base_capacity, 100)  # Floor to prevent division weirdness
        fill = total_words / base_capacity
        return min(1.0, max(0.5, fill)), actual_pages


# ── Section Header Detection ──────────────────────────────────────────

def _is_section_header(p) -> bool:
    """Check if a paragraph is a section header (all-bold, all-caps, short)."""
    text = p.text.strip()
    if not text or len(text) > 50:
        return False
    if text.upper() != text:
        return False
    if not p.runs:
        return False
    non_empty = [r for r in p.runs if r.text and r.text.strip()]
    return bool(non_empty) and all(r.bold for r in non_empty)


def _is_role_header(p) -> bool:
    """Check if a paragraph is a role/company header (has bold, contains date, has separator)."""
    text = p.text.strip()
    if not text or len(text) < 20:
        return False
    has_bold = any(r.bold for r in p.runs if r.text and r.text.strip()) if p.runs else False
    has_sep = "|" in text or "–" in text
    has_date = bool(re.search(r'(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+\d{4}', text))
    return has_bold and has_sep and has_date


def _is_bullet(p) -> bool:
    """Check if a paragraph is a bullet point."""
    text = p.text.strip()
    return text.startswith("•") or text.startswith("·")


# ── Adjustment Functions ──────────────────────────────────────────────

def _adjust_section_spacing(doc: Document, delta_pt: float) -> bool:
    """
    Adjust space_before on section headers by delta_pt.
    Positive = more space, Negative = less space.
    Returns True if any adjustment was made.
    """
    adjusted = False
    for p in doc.paragraphs:
        if not _is_section_header(p):
            continue

        pf = p.paragraph_format
        current = pf.space_before
        if current is None:
            current = Pt(8)  # Default assumption

        new_val = current + Pt(delta_pt)

        # Enforce bounds
        if new_val < SECTION_SPACING_MIN:
            new_val = SECTION_SPACING_MIN
        if new_val > SECTION_SPACING_MAX:
            new_val = SECTION_SPACING_MAX

        if new_val != current:
            pf.space_before = new_val
            adjusted = True

    return adjusted


def _adjust_paragraph_spacing(doc: Document, delta_pt: float) -> bool:
    """
    Adjust space_after on bullet paragraphs by delta_pt.
    Positive = more space, Negative = less space.
    Returns True if any adjustment was made.
    """
    adjusted = False
    for p in doc.paragraphs:
        if not _is_bullet(p):
            continue

        pf = p.paragraph_format
        current = pf.space_after
        if current is None:
            current = Pt(1)  # Default assumption for bullets

        new_val = current + Pt(delta_pt)

        # Enforce bounds
        if new_val < PARA_SPACING_MIN:
            new_val = PARA_SPACING_MIN
        if new_val > PARA_SPACING_MAX:
            new_val = PARA_SPACING_MAX

        if new_val != current:
            pf.space_after = new_val
            adjusted = True

    return adjusted


def _adjust_margins(doc: Document, delta_inches: float) -> bool:
    """
    Adjust all four margins by delta_inches.
    Positive = wider margins (less text area), Negative = narrower margins (more text area).
    Returns True if any adjustment was made.
    """
    adjusted = False
    for section in doc.sections:
        current = section.left_margin.inches  # All four are the same in our templates

        new_val = current + delta_inches

        # Enforce bounds
        if new_val < MARGIN_MIN:
            new_val = MARGIN_MIN
        if new_val > MARGIN_MAX:
            new_val = MARGIN_MAX

        if abs(new_val - current) > 0.001:
            section.top_margin = Inches(new_val)
            section.bottom_margin = Inches(new_val)
            section.left_margin = Inches(new_val)
            section.right_margin = Inches(new_val)
            adjusted = True

    return adjusted


def _adjust_body_font(doc: Document, delta_pt: float) -> bool:
    """
    Adjust body font size (runs that are 10pt ± 0.5pt) by delta_pt.
    Does NOT touch headers (20pt name, 11pt section headers, 9.5pt contact).
    Returns True if any adjustment was made.
    """
    adjusted = False
    for p in doc.paragraphs:
        for run in p.runs:
            if run.font.size is None:
                continue
            size_pt = run.font.size.pt
            # Only adjust body text (9.5-10.5pt range)
            if 9.0 <= size_pt <= 11.0 and size_pt not in (11.0, 20.0):
                new_size = Pt(size_pt + delta_pt)
                if new_size < FONT_MIN:
                    new_size = FONT_MIN
                if new_size > FONT_MAX:
                    new_size = FONT_MAX
                if new_size != run.font.size:
                    run.font.size = new_size
                    adjusted = True

    return adjusted


# ── Section Split Prevention ──────────────────────────────────────────

def _apply_keep_with_next(doc: Document) -> int:
    """
    Apply keepWithNext to section headers and role headers to prevent
    them from being orphaned at the bottom of a page.

    Returns the number of paragraphs modified.
    """
    count = 0
    for i, p in enumerate(doc.paragraphs):
        if _is_section_header(p) or _is_role_header(p):
            # Set keepWithNext via XML
            pPr = p._element.find(qn("w:pPr"))
            if pPr is None:
                pPr = etree.SubElement(p._element, qn("w:pPr"))

            # Remove existing keepNext if any
            for existing in pPr.findall(qn("w:keepNext")):
                pPr.remove(existing)

            # Add keepNext
            etree.SubElement(pPr, qn("w:keepNext"))
            count += 1

    return count


# ── Main Fitter ───────────────────────────────────────────────────────

def fit_resume(
    docx_path: str,
    target_pages: int,
    working_dir: Optional[str] = None,
) -> Dict:
    """
    Fit a resume to its target page count by adjusting visual properties.

    Args:
        docx_path: Path to the .docx resume file (modified in-place)
        target_pages: Expected page count (1 or 2)
        working_dir: Optional directory for temporary PDF conversions

    Returns:
        Dict with:
        - mode: "fill", "trim", or "none" (no adjustment needed)
        - adjustments: list of adjustments applied
        - iterations: number of measure-adjust cycles
        - initial_fill: starting fill ratio
        - final_fill: ending fill ratio
        - initial_pages: starting page count
        - final_pages: ending page count
    """
    # Use a temp working copy for PDF conversion to avoid clobbering
    if working_dir:
        work_path = os.path.join(working_dir, "fitter_work.docx")
    else:
        work_path = docx_path.rsplit(".", 1)[0] + "_fitter.docx"

    result = {
        "mode": "none",
        "adjustments": [],
        "iterations": 0,
        "section_splits_fixed": 0,
    }

    # Step 0: Apply keepWithNext to prevent section splits (always)
    doc = Document(docx_path)
    splits_fixed = _apply_keep_with_next(doc)
    doc.save(docx_path)
    result["section_splits_fixed"] = splits_fixed

    # Step 1: Measure initial fill
    initial_fill, initial_pages = _estimate_page_fill(docx_path, target_pages)
    result["initial_fill"] = round(initial_fill, 3)
    result["initial_pages"] = initial_pages

    # Step 2: Determine mode
    if initial_pages > target_pages:
        result["mode"] = "trim"
    elif initial_fill < FILL_LOW:
        result["mode"] = "fill"
    else:
        # Already in the sweet spot
        result["final_fill"] = round(initial_fill, 3)
        result["final_pages"] = initial_pages
        return result

    # Step 3: Apply adjustments iteratively
    if result["mode"] == "fill":
        # FILL mode: expand to fill whitespace
        fill_steps = [
            ("section_spacing", lambda d: _adjust_section_spacing(d, 1.5)),
            ("paragraph_spacing", lambda d: _adjust_paragraph_spacing(d, 0.5)),
            ("margins", lambda d: _adjust_margins(d, MARGIN_STEP)),
            ("body_font", lambda d: _adjust_body_font(d, 0.5)),
        ]
    else:
        # TRIM mode: compress to fit
        fill_steps = [
            ("paragraph_spacing", lambda d: _adjust_paragraph_spacing(d, -0.5)),
            ("section_spacing", lambda d: _adjust_section_spacing(d, -1.5)),
            ("margins", lambda d: _adjust_margins(d, -MARGIN_STEP)),
            ("body_font", lambda d: _adjust_body_font(d, -0.5)),
        ]

    step_idx = 0
    for iteration in range(MAX_ITERATIONS):
        if step_idx >= len(fill_steps):
            break

        step_name, step_fn = fill_steps[step_idx]

        doc = Document(docx_path)
        adjusted = step_fn(doc)

        if not adjusted:
            # This adjustment hit its bound — move to next
            step_idx += 1
            continue

        doc.save(docx_path)
        result["iterations"] += 1

        # Re-measure
        fill, pages = _estimate_page_fill(docx_path, target_pages)
        result["adjustments"].append({
            "step": step_name,
            "iteration": iteration + 1,
            "fill_after": round(fill, 3),
            "pages_after": pages,
        })

        # Check if we've achieved the goal
        if result["mode"] == "trim":
            if pages <= target_pages:
                break
        else:
            if fill >= FILL_LOW:
                break

            # If we're getting close, move to next adjustment type
            # to avoid overshooting with coarse adjustments
            if fill > FILL_LOW - 0.05:
                step_idx += 1

    # Final measurement
    final_fill, final_pages = _estimate_page_fill(docx_path, target_pages)
    result["final_fill"] = round(final_fill, 3)
    result["final_pages"] = final_pages

    # Safety check: if we overshot (went over pages), undo last step
    if final_pages > target_pages and result["mode"] == "fill":
        # Revert the last adjustment by trimming
        doc = Document(docx_path)
        if result["adjustments"]:
            last = result["adjustments"][-1]["step"]
            if last == "margins":
                _adjust_margins(doc, -MARGIN_STEP)
            elif last == "body_font":
                _adjust_body_font(doc, -0.5)
            elif last == "section_spacing":
                _adjust_section_spacing(doc, -1.5)
            elif last == "paragraph_spacing":
                _adjust_paragraph_spacing(doc, -0.5)
            doc.save(docx_path)
            result["adjustments"].append({"step": f"revert_{last}", "reason": "overflow"})

        final_fill, final_pages = _estimate_page_fill(docx_path, target_pages)
        result["final_fill"] = round(final_fill, 3)
        result["final_pages"] = final_pages

    # Clean up temp PDF if any
    pdf_cleanup = docx_path.rsplit(".", 1)[0] + ".pdf"
    if os.path.exists(pdf_cleanup):
        try:
            os.remove(pdf_cleanup)
        except Exception:
            logger.debug("Could not clean up temp PDF: %s", pdf_cleanup)

    return result


def format_fit_report(result: Dict) -> str:
    """Format a fit result as a human-readable summary."""
    lines = []
    lines.append(f"Resume Fitter: {result['mode'].upper()} mode")
    lines.append(f"  Section splits fixed: {result['section_splits_fixed']}")
    lines.append(f"  Initial: {result['initial_fill']:.0%} fill, {result.get('initial_pages', '?')} pages")
    lines.append(f"  Final:   {result.get('final_fill', '?'):.0%} fill, {result.get('final_pages', '?')} pages")

    if result["adjustments"]:
        lines.append(f"  Adjustments ({result['iterations']} iterations):")
        for adj in result["adjustments"]:
            if "reason" in adj:
                lines.append(f"    - {adj['step']} ({adj['reason']})")
            else:
                lines.append(f"    - {adj['step']}: {adj['fill_after']:.0%} fill, {adj['pages_after']} pages")

    return "\n".join(lines)
