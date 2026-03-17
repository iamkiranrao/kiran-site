"""
Resume Pipeline Orchestrator — 12-step async pipeline with SSE progress events.

Supports two modes:
A) Full pipeline (original): run_pipeline() — runs all 12 steps in one go
B) Phased workflow (new):
   Phase 1: run_analysis_phase() — Analyze JD → Research → Generate proposal → Pause for review
   Phase 2: run_generation_phase() — Apply approved edits → Generate docs → Package

Takes a JD and produces 5 documents:
1. Customized Resume (.docx)
2. Match Score & Gap Analysis (.docx)
3. Cover Letter (.docx)
4. Company Research Brief (.docx)
5. Interview Question Bank (.docx)
"""

import os
import json
import uuid
import shutil
import zipfile
import asyncio
import logging
from typing import AsyncGenerator, List, Optional
from docx import Document

logger = logging.getLogger(__name__)

from services.resume_editor import (
    pt, rewrite, swap, has_br, get_lines, rewrite_br_paragraph,
    bold_sweep, detect_persona, copy_template, read_template_structure,
    verify_resume, VerbTracker, find_summary, find_tagline, reorder_skills,
    ensure_not_bold, find_bullet_paragraphs, get_experience_bullet_map,
    extract_template_content, write_core_competencies, write_technical_skills,
)
from services.skills_library import select_skills
from services.claude_client import (
    analyze_jd, generate_resume_content, generate_match_score,
    generate_cover_letter, generate_company_brief, generate_interview_questions,
    generate_strategy_proposal, refine_resume_content, refine_section, discuss_section,
    generate_pre_match_score, audit_resume_content, extract_ats_keywords,
)
from services.doc_creator import (
    create_match_score, create_cover_letter,
    create_company_brief, create_interview_questions,
)
from services.resume_scorer import score_resume, format_scorecard, compare_scores
from services.resume_fitter import fit_resume
from services.resume_certifier import certify_resume


from utils.config import data_dir
JOBS_DIR = data_dir("jobs")

PIPELINE_STEPS = [
    {"step": 1, "label": "Parsing inputs"},
    {"step": 2, "label": "Analyzing job description"},
    {"step": 3, "label": "Selecting template"},
    {"step": 4, "label": "Reading template structure"},
    {"step": 5, "label": "Generating customized content"},
    {"step": 6, "label": "Applying edits to resume"},
    {"step": 7, "label": "Running quality checks"},
    {"step": 8, "label": "Generating match score"},
    {"step": 9, "label": "Writing cover letter"},
    {"step": 10, "label": "Compiling company brief"},
    {"step": 11, "label": "Building interview questions"},
    {"step": 12, "label": "Packaging deliverables"},
]

# Analysis phase uses different step labels
ANALYSIS_STEPS = [
    {"step": 1, "label": "Parsing inputs"},
    {"step": 2, "label": "Analyzing job description"},
    {"step": 3, "label": "Selecting template"},
    {"step": 4, "label": "Reading template structure"},
    {"step": 5, "label": "Researching & building strategy"},
    {"step": 6, "label": "Generating proposal"},
]

# Generation phase steps (after approval)
GENERATION_STEPS = [
    {"step": 1, "label": "Loading approved proposal"},
    {"step": 2, "label": "Applying edits to resume"},
    {"step": 3, "label": "Running quality checks"},
    {"step": 4, "label": "Auditing content rules"},
    {"step": 5, "label": "Analyzing ATS keywords"},
    {"step": 6, "label": "Generating match score"},
    {"step": 7, "label": "Writing cover letter"},
    {"step": 8, "label": "Compiling company brief"},
    {"step": 9, "label": "Building interview questions"},
    {"step": 10, "label": "Packaging deliverables"},
]


def _emit(step: int, status: str, detail: str = "", total_steps: int = 12, steps_def: list = None) -> str:
    """Create an SSE event JSON string."""
    steps = steps_def or PIPELINE_STEPS
    label = steps[step - 1]["label"] if step <= len(steps) else "Unknown"
    event = {
        "step": step,
        "total_steps": total_steps,
        "label": label,
        "status": status,
        "detail": detail,
    }
    return json.dumps(event)


def _emit_analysis(step: int, status: str, detail: str = "") -> str:
    return _emit(step, status, detail, total_steps=6, steps_def=ANALYSIS_STEPS)


def _emit_generation(step: int, status: str, detail: str = "") -> str:
    return _emit(step, status, detail, total_steps=10, steps_def=GENERATION_STEPS)


# ═══════════════════════════════════════════════════════════════════════════════
# Shared: Bullet count validation
# ═══════════════════════════════════════════════════════════════════════════════

def _validate_bullet_count(revised, expected_count: int, fallback):
    """Validate that revised experience bullets match expected count.
    If count is wrong, pad with fallback bullets or truncate to fit."""
    if not isinstance(revised, list) or expected_count is None:
        return revised
    if len(revised) == expected_count:
        return revised
    if len(revised) > expected_count:
        # Truncate
        return revised[:expected_count]
    # Pad with fallback bullets if available
    if isinstance(fallback, list):
        while len(revised) < expected_count and len(fallback) > len(revised):
            revised.append(fallback[len(revised)])
    return revised


# ═══════════════════════════════════════════════════════════════════════════════
# Shared: Apply content to resume document
# ═══════════════════════════════════════════════════════════════════════════════

def _apply_content_to_resume(doc: Document, content: dict, version: str, jd_text: str = None) -> None:
    """Apply generated/approved content to a resume document. Shared by both pipeline modes."""
    verb_tracker = VerbTracker()

    # Rewrite tagline
    if "tagline" in content and isinstance(content.get("tagline"), str):
        tagline_idx, tagline_p = find_tagline(doc)
        if tagline_p and content.get("tagline"):
            rewrite(tagline_p, content["tagline"])

    # Rewrite summary
    if "summary" in content and isinstance(content.get("summary"), str):
        summary_idx, summary_p = find_summary(doc)
        if summary_p and content.get("summary"):
            rewrite(summary_p, content["summary"])

    # Rewrite career highlights
    if "career_highlights" in content and isinstance(content["career_highlights"], list):
        for i, p in enumerate(doc.paragraphs):
            text = pt(p).strip()
            if text == "CAREER HIGHLIGHTS":
                for j in range(i + 1, min(i + 3, len(doc.paragraphs))):
                    if has_br(doc.paragraphs[j]):
                        new_lines = []
                        for bullet in content["career_highlights"]:
                            if not bullet.startswith("•"):
                                bullet = f"• {bullet}"
                            first_word = bullet.lstrip("• ").split()[0] if bullet.lstrip("• ").split() else ""
                            verb_tracker.register(first_word)
                            new_lines.append(bullet)
                        rewrite_br_paragraph(doc.paragraphs[j], new_lines)
                        break
                break

    # Rewrite experience bullets
    if "experience_bullets" in content and isinstance(content["experience_bullets"], dict):
        for company_name, bullets in content["experience_bullets"].items():
            header_idx, bullet_indices = find_bullet_paragraphs(doc, company_name)

            if header_idx is None or not bullet_indices:
                continue

            new_bullets = []
            for bullet in bullets:
                if not bullet.startswith("•"):
                    bullet = f"• {bullet}"
                words = bullet.lstrip("• ").split()
                if words:
                    verb = words[0]
                    unique_verb = verb_tracker.pick_verb(verb)
                    if unique_verb != verb:
                        bullet = bullet.replace(verb, unique_verb, 1)
                new_bullets.append(bullet)

            if len(bullet_indices) == 1 and has_br(doc.paragraphs[bullet_indices[0]]):
                rewrite_br_paragraph(doc.paragraphs[bullet_indices[0]], new_bullets)
            else:
                for k, bidx in enumerate(bullet_indices):
                    if k < len(new_bullets):
                        rewrite(doc.paragraphs[bidx], new_bullets[k])

    # Write skills from locked-in proposal content (set during analysis by the
    # skills library, then potentially refined by the user during review).
    skills_written = False

    # Try core competencies from proposal
    core_skills = content.get("core_competencies_priority")
    if core_skills and isinstance(core_skills, list):
        try:
            skills_written = write_core_competencies(doc, core_skills) or skills_written
        except Exception:
            logger.exception("Failed to write skills to template")

    # Try technical skills from proposal
    tech_skills = content.get("technical_skills_priority")
    if tech_skills:
        try:
            skills_written = write_technical_skills(doc, tech_skills, version) or skills_written
        except Exception:
            logger.warning("Try technical skills from proposal failed (non-blocking)", exc_info=True)

    # Fallback: run skills library directly if proposal didn't have the fields
    if not skills_written and jd_text:
        try:
            skills_result = select_skills(jd_text, version=version)
            core_written = write_core_competencies(doc, skills_result["core_competencies"])
            tech_written = write_technical_skills(doc, skills_result["technical_skills"], version)
            skills_written = core_written or tech_written
        except Exception:
            skills_written = False

    # Last resort: simple reordering from legacy skills_priority field
    if not skills_written:
        fallback_skills = content.get("skills_priority")
        if fallback_skills and isinstance(fallback_skills, list):
            found_skills = False
            for i, p in enumerate(doc.paragraphs):
                text = pt(p).strip()
                if text.upper() in ("SKILLS", "CORE SKILLS", "TECHNICAL SKILLS",
                                    "LEADERSHIP & STRATEGIC COMPETENCIES",
                                    "CORE COMPETENCIES", "TECHNICAL COMPETENCIES"):
                    for j in range(i + 1, min(i + 8, len(doc.paragraphs))):
                        skills_p = doc.paragraphs[j]
                        skills_text = pt(skills_p).strip()
                        if not skills_text:
                            break
                        s_all_bold = all(r.bold for r in skills_p.runs if r.text and r.text.strip()) if skills_p.runs else False
                        if s_all_bold and len(skills_text) < 40:
                            break
                        if "," in skills_text or "│" in skills_text or "|" in skills_text:
                            reorder_skills(skills_p, fallback_skills)
                            found_skills = True
                    if found_skills:
                        break

    # Bold sweep
    bold_sweep(doc)


# ═══════════════════════════════════════════════════════════════════════════════
# Proposal persistence (JSON files in job dir)
# ═══════════════════════════════════════════════════════════════════════════════

def _save_proposal(job_id: str, data: dict) -> None:
    """Save proposal + metadata to job dir."""
    job_dir = os.path.join(JOBS_DIR, job_id)
    path = os.path.join(job_dir, "_proposal.json")
    with open(path, "w") as f:
        json.dump(data, f, indent=2)


def _load_proposal(job_id: str) -> Optional[dict]:
    """Load saved proposal from job dir."""
    path = os.path.join(JOBS_DIR, job_id, "_proposal.json")
    if not os.path.exists(path):
        return None
    with open(path) as f:
        return json.load(f)


def get_proposal(job_id: str) -> Optional[dict]:
    """Public accessor for the proposal endpoint."""
    return _load_proposal(job_id)


# ═══════════════════════════════════════════════════════════════════════════════
# Application Tracker (persists across runs)
# ═══════════════════════════════════════════════════════════════════════════════

TRACKER_PATH = os.path.join(data_dir("jobs"), "application_tracker.json")


def _load_tracker() -> list:
    """Load the application tracker index."""
    if os.path.exists(TRACKER_PATH):
        with open(TRACKER_PATH) as f:
            return json.load(f)
    return []


def _save_tracker(entries: list) -> None:
    """Save the application tracker index."""
    os.makedirs(os.path.dirname(TRACKER_PATH), exist_ok=True)
    with open(TRACKER_PATH, "w") as f:
        json.dump(entries, f, indent=2)


def add_tracker_entry(entry: dict) -> None:
    """Add a new entry to the application tracker."""
    entries = _load_tracker()
    entries.append(entry)
    _save_tracker(entries)


def get_tracker() -> list:
    """Get all tracker entries."""
    return _load_tracker()


# ═══════════════════════════════════════════════════════════════════════════════
# Parallel doc generation (during review phase)
# ═══════════════════════════════════════════════════════════════════════════════

async def start_parallel_doc_generation(job_id: str, api_key: str) -> dict:
    """
    Kick off cover letter, company brief, and interview questions in parallel.
    These don't depend on the final resume text — only on JD analysis + background.
    Saves results to job dir so generation phase can skip those steps.
    """
    saved = _load_proposal(job_id)
    if not saved:
        return {"error": f"No proposal found for job {job_id}"}

    jd_text = saved["jd_text"]
    jd_analysis = saved["jd_analysis"]
    company = saved["company"]
    role = saved["role"]
    job_dir = os.path.join(JOBS_DIR, job_id)

    # Run all three in parallel
    letter_task = generate_cover_letter(api_key, jd_text, jd_analysis)
    brief_task = generate_company_brief(api_key, jd_analysis)
    questions_task = generate_interview_questions(api_key, jd_text, jd_analysis)

    letter_content, brief_content, questions_content = await asyncio.gather(
        letter_task, brief_task, questions_task
    )

    # Create the docs and save to job dir
    letter_path = os.path.join(job_dir, "Cover_Letter.docx")
    create_cover_letter(letter_content, company, role, letter_path)

    brief_path = os.path.join(job_dir, "Company_Brief.docx")
    create_company_brief(brief_content, company, role, brief_path)

    questions_path = os.path.join(job_dir, "Interview_Questions.docx")
    create_interview_questions(questions_content, company, role, questions_path)

    # Mark as pre-generated in proposal
    saved["parallel_docs_ready"] = True
    _save_proposal(job_id, saved)

    return {"status": "completed", "files": ["Cover_Letter.docx", "Company_Brief.docx", "Interview_Questions.docx"]}


def update_proposal_feedback(job_id: str, feedback: str) -> None:
    """Save user feedback into the existing proposal."""
    proposal = _load_proposal(job_id)
    if proposal:
        proposal["user_feedback"] = feedback
        _save_proposal(job_id, proposal)


async def refine_proposal_section(
    job_id: str,
    section_id: str,
    feedback: str,
    api_key: str,
) -> dict:
    """Refine a single section of the proposal based on user feedback."""
    saved = _load_proposal(job_id)
    if not saved:
        return {"error": f"No proposal found for job {job_id}"}

    jd_analysis = saved["jd_analysis"]
    proposal = saved["proposal"]
    original_content = saved.get("original_content", {})
    proposed_changes = proposal.get("proposed_changes", {})

    # Find the original and proposed text for this section
    original_text = None
    current_proposed = None
    bullet_count = None

    for section in original_content.get("sections", []):
        if section["id"] == section_id:
            original_text = section["current_text"]
            break

    # Map section_id to the right field in proposed_changes
    if section_id == "tagline":
        current_proposed = proposed_changes.get("tagline", "")
    elif section_id == "summary":
        current_proposed = proposed_changes.get("summary", "")
    elif section_id == "career_highlights":
        current_proposed = proposed_changes.get("career_highlights", [])
    elif section_id == "core_competencies":
        current_proposed = proposed_changes.get("core_competencies_priority", proposed_changes.get("skills_priority", []))
    elif section_id == "technical_competencies":
        current_proposed = proposed_changes.get("technical_skills_priority", proposed_changes.get("skills_priority", []))
    elif section_id == "skills":  # legacy fallback
        current_proposed = proposed_changes.get("skills_priority", [])
    elif section_id.startswith("exp__"):
        company_key = section_id.replace("exp__", "")
        exp_bullets = proposed_changes.get("experience_bullets", {})
        # Find matching company key (case-insensitive partial match)
        for key, bullets in exp_bullets.items():
            if company_key.lower() in key.lower() or key.lower() in company_key.lower():
                current_proposed = bullets
                bullet_count = len(bullets)
                break

    if current_proposed is None:
        return {"error": f"Section {section_id} not found in proposal"}

    result = await refine_section(
        api_key=api_key,
        section_id=section_id,
        section_label=section_id.replace("exp__", ""),
        original_text=original_text or current_proposed,
        current_proposed_text=current_proposed,
        feedback=feedback,
        jd_analysis=jd_analysis,
        bullet_count=bullet_count,
    )

    # Update the proposal with revised text
    revised = result.get("revised_text", current_proposed)

    # Validate bullet count for experience sections
    if section_id.startswith("exp__") and bullet_count:
        revised = _validate_bullet_count(revised, bullet_count, current_proposed)

    if section_id == "tagline":
        proposed_changes["tagline"] = revised
    elif section_id == "summary":
        proposed_changes["summary"] = revised
    elif section_id == "career_highlights":
        proposed_changes["career_highlights"] = revised
    elif section_id == "core_competencies":
        proposed_changes["core_competencies_priority"] = revised
    elif section_id == "technical_competencies":
        proposed_changes["technical_skills_priority"] = revised
    elif section_id == "skills":  # legacy fallback
        proposed_changes["skills_priority"] = revised
    elif section_id.startswith("exp__"):
        company_key = section_id.replace("exp__", "")
        exp_bullets = proposed_changes.get("experience_bullets", {})
        for key in exp_bullets:
            if company_key.lower() in key.lower() or key.lower() in company_key.lower():
                exp_bullets[key] = revised
                break

    saved["proposal"]["proposed_changes"] = proposed_changes
    _save_proposal(job_id, saved)

    return {
        "section_id": section_id,
        "revised_text": revised,
        "rationale": result.get("rationale", ""),
    }


async def discuss_proposal_section(
    job_id: str,
    section_id: str,
    message: str,
    conversation_history: list,
    api_key: str,
) -> dict:
    """Conversational review of a section — questions, rationale, AND changes."""
    saved = _load_proposal(job_id)
    if not saved:
        return {"error": f"No proposal found for job {job_id}"}

    jd_analysis = saved["jd_analysis"]
    proposal = saved["proposal"]
    original_content = saved.get("original_content", {})
    proposed_changes = proposal.get("proposed_changes", {})

    # Find the original and proposed text for this section
    original_text = None
    current_proposed = None
    bullet_count = None

    for section in original_content.get("sections", []):
        if section["id"] == section_id:
            original_text = section["current_text"]
            break

    # Map section_id to the right field in proposed_changes
    if section_id == "tagline":
        current_proposed = proposed_changes.get("tagline", "")
    elif section_id == "summary":
        current_proposed = proposed_changes.get("summary", "")
    elif section_id == "career_highlights":
        current_proposed = proposed_changes.get("career_highlights", [])
    elif section_id == "core_competencies":
        current_proposed = proposed_changes.get("core_competencies_priority", proposed_changes.get("skills_priority", []))
    elif section_id == "technical_competencies":
        current_proposed = proposed_changes.get("technical_skills_priority", proposed_changes.get("skills_priority", []))
    elif section_id == "skills":  # legacy fallback
        current_proposed = proposed_changes.get("skills_priority", [])
    elif section_id.startswith("exp__"):
        company_key = section_id.replace("exp__", "")
        exp_bullets = proposed_changes.get("experience_bullets", {})
        for key, bullets in exp_bullets.items():
            if company_key.lower() in key.lower() or key.lower() in company_key.lower():
                current_proposed = bullets
                bullet_count = len(bullets)
                break

    if current_proposed is None:
        return {"error": f"Section {section_id} not found in proposal"}

    result = await discuss_section(
        api_key=api_key,
        section_id=section_id,
        section_label=section_id.replace("exp__", ""),
        original_text=original_text or current_proposed,
        current_proposed_text=current_proposed,
        message_text=message,
        conversation_history=conversation_history,
        jd_analysis=jd_analysis,
        bullet_count=bullet_count,
    )

    # If changes were made, update the proposal
    if result.get("has_changes") and result.get("revised_text") is not None:
        revised = result["revised_text"]

        # Validate bullet count for experience sections
        if section_id.startswith("exp__") and bullet_count:
            revised = _validate_bullet_count(revised, bullet_count, current_proposed)

        if section_id == "tagline":
            proposed_changes["tagline"] = revised
        elif section_id == "summary":
            proposed_changes["summary"] = revised
        elif section_id == "career_highlights":
            proposed_changes["career_highlights"] = revised
        elif section_id == "core_competencies":
            proposed_changes["core_competencies_priority"] = revised
        elif section_id == "technical_competencies":
            proposed_changes["technical_skills_priority"] = revised
        elif section_id == "skills":  # legacy fallback
            proposed_changes["skills_priority"] = revised
        elif section_id.startswith("exp__"):
            company_key = section_id.replace("exp__", "")
            exp_bullets = proposed_changes.get("experience_bullets", {})
            for key in exp_bullets:
                if company_key.lower() in key.lower() or key.lower() in company_key.lower():
                    exp_bullets[key] = revised
                    break

        saved["proposal"]["proposed_changes"] = proposed_changes
        _save_proposal(job_id, saved)

    return {
        "section_id": section_id,
        "message": result.get("message", ""),
        "revised_text": result.get("revised_text"),
        "has_changes": result.get("has_changes", False),
    }


def lock_section(job_id: str, section_id: str) -> dict:
    """Mark a section as approved/locked."""
    saved = _load_proposal(job_id)
    if not saved:
        return {"error": f"No proposal found for job {job_id}"}

    saved.setdefault("approved_sections", {})[section_id] = True
    _save_proposal(job_id, saved)

    total = len(saved.get("original_content", {}).get("sections", []))
    locked = len(saved.get("approved_sections", {}))

    return {
        "section_id": section_id,
        "locked": True,
        "total_sections": total,
        "locked_sections": locked,
        "all_locked": locked >= total,
    }


# ═══════════════════════════════════════════════════════════════════════════════
# Phase 1: Analysis + Proposal
# ═══════════════════════════════════════════════════════════════════════════════

async def run_analysis_phase(
    jd_text: str,
    persona: str,
    version: str,
    api_key: str,
) -> AsyncGenerator[str, None]:
    """
    Phase 1: Analyze JD → Select template → Read structure → Generate strategy proposal.
    Saves state to disk so Phase 2 can resume.
    """
    job_id = str(uuid.uuid4())[:8]
    job_dir = os.path.join(JOBS_DIR, job_id)
    os.makedirs(job_dir, mode=0o755, exist_ok=True)

    try:
        # ── Step 1: Parse inputs ─────────────────────────────────────────
        yield _emit_analysis(1, "in_progress")

        if persona == "auto":
            persona = detect_persona(jd_text)

        yield _emit_analysis(1, "completed", f"Persona: {persona}, Version: {version}")

        # ── Step 2: Analyze JD ───────────────────────────────────────────
        yield _emit_analysis(2, "in_progress")

        jd_analysis = await analyze_jd(api_key, jd_text)

        company = jd_analysis.get("company", "Unknown")
        role = jd_analysis.get("job_title", "Product Manager")
        yield _emit_analysis(2, "completed", f"{role} at {company}")

        # ── Step 3: Select and copy template ─────────────────────────────
        yield _emit_analysis(3, "in_progress")

        resume_path = os.path.join(job_dir, "Resume.docx")
        copy_template(persona, version, resume_path)

        yield _emit_analysis(3, "completed", f"{persona} {version} template")

        # ── Step 4: Read template structure ──────────────────────────────
        yield _emit_analysis(4, "in_progress")

        structure = read_template_structure(resume_path)
        template_doc = Document(resume_path)
        bullet_map = get_experience_bullet_map(template_doc)
        structure["experience_entries"] = bullet_map

        # Extract original text from each section for before/after display
        original_content = extract_template_content(template_doc)

        yield _emit_analysis(4, "completed", f"{structure['total_paragraphs']} paragraphs, {len(structure['sections'])} sections")

        # Score baseline template (before customization)
        baseline_score = None
        try:
            baseline_score = score_resume(resume_path)
        except Exception:
            logger.warning("Baseline scoring failed (non-blocking)", exc_info=True)

        # ── Step 5: Research & build strategy ────────────────────────────
        yield _emit_analysis(5, "in_progress", "Researching company and building customization strategy...")

        proposal = await generate_strategy_proposal(
            api_key, jd_text, jd_analysis, structure, persona, version
        )

        if "error" in proposal:
            yield _emit_analysis(5, "warning", f"Partial result: {proposal.get('error', '')}")
        else:
            yield _emit_analysis(5, "completed", "Strategy and proposed changes ready")

        # ── Step 5b: Run skills library and inject into proposal ──────────
        # The skills library does domain-optimized selection; its output
        # replaces whatever Claude proposed for core_competencies and
        # technical_skills so the review flow shows the actual selections.
        try:
            skills_result = select_skills(jd_text, version=version)
            proposed = proposal.get("proposed_changes", {})
            proposed["core_competencies_priority"] = skills_result["core_competencies"]
            # Technical skills: flatten categorized dict for review display
            tech = skills_result["technical_skills"]
            if isinstance(tech, dict):
                proposed["technical_skills_priority"] = tech  # keep as dict for categorized display
            else:
                proposed["technical_skills_priority"] = tech
            proposed["core_competencies_rationale"] = (
                f"Domain-optimized selection from skills library. "
                f"Detected domains: {', '.join(skills_result.get('detected_domains', []))}. "
                + proposed.get("core_competencies_rationale", "")
            )
            proposed["technical_skills_rationale"] = (
                f"Domain-optimized selection from skills library. "
                f"Detected domains: {', '.join(skills_result.get('detected_domains', []))}. "
                + proposed.get("technical_skills_rationale", "")
            )
            proposal["proposed_changes"] = proposed
        except Exception:
            logger.warning("Skills library proposal failed, using Claude proposal", exc_info=True)

        # Generate pre-match score in background (non-blocking best effort)
        pre_match_score = 60
        try:
            pre_match_score = await generate_pre_match_score(api_key, jd_text, jd_analysis)
        except Exception:
            logger.warning("Pre-match score generation failed (non-blocking)", exc_info=True)

        # ── Step 6: Save proposal for review ─────────────────────────────
        yield _emit_analysis(6, "in_progress")

        # Save everything needed for Phase 2
        saved_state = {
            "job_id": job_id,
            "jd_text": jd_text,
            "jd_analysis": jd_analysis,
            "persona": persona,
            "version": version,
            "company": company,
            "role": role,
            "template_structure": structure,
            "original_content": original_content,
            "proposal": proposal,
            "pre_match_score": pre_match_score,
            "is_ic_role": jd_analysis.get("is_ic_role", False),
            "ic_signals": jd_analysis.get("ic_signals", []),
            "baseline_score": baseline_score,
            "approved_sections": {},  # Track which sections user has locked in
            "status": "awaiting_review",
        }
        _save_proposal(job_id, saved_state)

        yield _emit_analysis(6, "completed", "Proposal saved — ready for review")

        # Final event with proposal data + original content for before/after
        final = {
            "step": 6,
            "total_steps": 6,
            "label": "Proposal Ready",
            "status": "completed",
            "job_id": job_id,
            "company": company,
            "role": role,
            "persona": persona,
            "version": version,
            "phase": "analysis",
            "strategy": proposal.get("strategy", ""),
            "proposed_changes": proposal.get("proposed_changes", {}),
            "original_content": original_content,
            "pre_match_score": pre_match_score,
            "is_ic_role": jd_analysis.get("is_ic_role", False),
            "ic_signals": jd_analysis.get("ic_signals", []),
            "baseline_score": baseline_score,
        }
        yield json.dumps(final)

    except Exception as e:
        yield json.dumps({
            "step": 0,
            "total_steps": 6,
            "status": "error",
            "label": "Analysis error",
            "detail": str(e),
        })


# ═══════════════════════════════════════════════════════════════════════════════
# Phase 2: Generation (after approval)
# ═══════════════════════════════════════════════════════════════════════════════

async def run_generation_phase(
    job_id: str,
    api_key: str,
    feedback: Optional[str] = None,
) -> AsyncGenerator[str, None]:
    """
    Phase 2: Load approved proposal → Apply edits → Generate docs → Package.
    """
    try:
        # ── Step 1: Load proposal ────────────────────────────────────────
        yield _emit_generation(1, "in_progress")

        saved = _load_proposal(job_id)
        if not saved:
            yield json.dumps({
                "step": 0, "total_steps": 10, "status": "error",
                "label": "Error", "detail": f"No proposal found for job {job_id}",
            })
            return

        job_dir = os.path.join(JOBS_DIR, job_id)
        jd_text = saved["jd_text"]
        jd_analysis = saved["jd_analysis"]
        persona = saved["persona"]
        version = saved["version"]
        company = saved["company"]
        role = saved["role"]
        structure = saved["template_structure"]
        proposal = saved["proposal"]
        resume_path = os.path.join(job_dir, "Resume.docx")

        # If feedback provided, refine the content
        if feedback and feedback.strip():
            yield _emit_generation(1, "in_progress", "Refining based on your feedback...")
            content = await refine_resume_content(
                api_key, proposal, feedback, jd_analysis, structure, persona, version
            )
            if "error" in content:
                yield _emit_generation(1, "warning", f"Used original proposal: {content.get('error', '')}")
                content = proposal.get("proposed_changes", {})
        else:
            content = proposal.get("proposed_changes", {})

        yield _emit_generation(1, "completed", "Proposal loaded")

        # ── Step 2: Apply edits to resume ────────────────────────────────
        yield _emit_generation(2, "in_progress")

        doc = Document(resume_path)
        _apply_content_to_resume(doc, content, version, jd_text=jd_text)
        doc.save(resume_path)

        yield _emit_generation(2, "completed", "Resume edits applied")

        # ── Step 3: Quality checks ───────────────────────────────────────
        yield _emit_generation(3, "in_progress")

        for attempt in range(3):
            result = verify_resume(resume_path)
            if result["passed"]:
                yield _emit_generation(3, "completed", f"Passed (attempt {attempt + 1})")
                break
            if attempt < 2:
                doc = Document(resume_path)
                bold_sweep(doc)
                doc.save(resume_path)
                yield _emit_generation(3, "in_progress", f"Auto-fixing (attempt {attempt + 2})...")
            else:
                detail = f"{result['error_count']} errors, {result['warning_count']} warnings"
                yield _emit_generation(3, "warning", detail)

        # Visual fitting: adjust spacing/margins/font to optimize page fill
        fit_result = None
        try:
            target_pages = 1 if version == "1-Page" else 2
            fit_result = fit_resume(resume_path, target_pages=target_pages)
            if fit_result["mode"] != "none":
                yield _emit_generation(3, "completed",
                    f"Visual fit: {fit_result['mode']} mode, "
                    f"{fit_result['initial_fill']:.0%} → {fit_result.get('final_fill', 0):.0%} fill")
        except Exception:
            logger.warning("Resume fitting failed (non-blocking)", exc_info=True)

        # Score customized resume (after edits, with JD for alignment)
        baseline_score = saved.get("baseline_score")
        post_score = None
        score_delta = None
        try:
            post_score = score_resume(resume_path, jd_text=jd_text)
            if baseline_score:
                score_delta = compare_scores(baseline_score, post_score)
        except Exception:
            logger.warning("Post-scoring comparison failed (non-blocking)", exc_info=True)

        # Extract resume text for audit and ATS checks
        resume_doc = Document(resume_path)
        resume_full_text = "\n".join(p.text for p in resume_doc.paragraphs if p.text.strip())

        # ── Step 4: Audit content rules ──────────────────────────────────
        yield _emit_generation(4, "in_progress")

        audit_result = await audit_resume_content(api_key, resume_full_text, jd_analysis)
        audit_detail = f"Score: {audit_result.get('score', '?')}/100"
        if audit_result.get("violations"):
            audit_detail += f", {len(audit_result['violations'])} violations"
        if audit_result.get("warnings"):
            audit_detail += f", {len(audit_result['warnings'])} warnings"

        yield _emit_generation(4, "completed" if audit_result.get("passed") else "warning", audit_detail)

        # ── Step 5: ATS keyword analysis ─────────────────────────────────
        yield _emit_generation(5, "in_progress")

        ats_result = await extract_ats_keywords(api_key, jd_analysis, resume_full_text)
        ats_detail = f"{ats_result.get('matched_count', 0)}/{ats_result.get('total_keywords', 0)} keywords ({ats_result.get('coverage_pct', 0)}%)"
        if ats_result.get("critical_missing"):
            ats_detail += f", {len(ats_result['critical_missing'])} critical gaps"

        yield _emit_generation(5, "completed", ats_detail)

        # ── Step 6: Match Score ──────────────────────────────────────────
        yield _emit_generation(6, "in_progress")

        match_content = await generate_match_score(api_key, jd_text, jd_analysis)
        match_path = os.path.join(job_dir, "Match_Score.docx")
        create_match_score(match_content, company, role, match_path)

        # Extract the numeric post-match score from the markdown
        post_match_score = 0
        import re
        score_match = re.search(r'Match Score:\s*(\d+)', match_content)
        if score_match:
            post_match_score = int(score_match.group(1))

        yield _emit_generation(6, "completed", f"Post-customization score: {post_match_score}/100")

        # ── Step 7-9: Cover Letter, Brief, Interview Questions ───────────
        # Check if parallel generation already produced these docs
        parallel_ready = saved.get("parallel_docs_ready", False)

        if parallel_ready and os.path.exists(os.path.join(job_dir, "Cover_Letter.docx")):
            yield _emit_generation(7, "completed", "Pre-generated during review")
            yield _emit_generation(8, "completed", "Pre-generated during review")
            yield _emit_generation(9, "completed", "Pre-generated during review")
        else:
            # Generate sequentially if not pre-generated
            yield _emit_generation(7, "in_progress")
            letter_content = await generate_cover_letter(api_key, jd_text, jd_analysis)
            letter_path = os.path.join(job_dir, "Cover_Letter.docx")
            create_cover_letter(letter_content, company, role, letter_path)
            yield _emit_generation(7, "completed")

            yield _emit_generation(8, "in_progress")
            brief_content = await generate_company_brief(api_key, jd_analysis)
            brief_path = os.path.join(job_dir, "Company_Brief.docx")
            create_company_brief(brief_content, company, role, brief_path)
            yield _emit_generation(8, "completed")

            yield _emit_generation(9, "in_progress")
            questions_content = await generate_interview_questions(api_key, jd_text, jd_analysis)
            questions_path = os.path.join(job_dir, "Interview_Questions.docx")
            create_interview_questions(questions_content, company, role, questions_path)
            yield _emit_generation(9, "completed")

        # ── Final certification ────────────────────────────────────────────
        target_pages = 1 if version == "1-Page" else 2
        certification = None
        try:
            certification = certify_resume(
                resume_path,
                target_pages=target_pages,
                jd_text=jd_text,
                scorer_result=post_score,
                verifier_result=result,  # from verify_resume loop
                fitter_result=fit_result,
            )
        except Exception:
            logger.warning("Certification failed (non-blocking)", exc_info=True)

        # ── Step 10: Package deliverables ─────────────────────────────────
        yield _emit_generation(10, "in_progress")

        zip_path = os.path.join(job_dir, "Application_Package.zip")
        files_to_zip = [
            "Resume.docx", "Match_Score.docx", "Cover_Letter.docx",
            "Company_Brief.docx", "Interview_Questions.docx",
        ]
        with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
            for fname in files_to_zip:
                fpath = os.path.join(job_dir, fname)
                if os.path.exists(fpath):
                    zf.write(fpath, fname)

        # Update proposal status
        saved["status"] = "completed"
        _save_proposal(job_id, saved)

        # ── Application tracker entry ─────────────────────────────────────
        from datetime import datetime
        # Review metadata: how many sections were discussed/refined
        approved_sections = saved.get("approved_sections", {})
        total_sections = len(saved.get("original_content", {}).get("sections", []))
        user_feedback_text = saved.get("user_feedback", None)
        is_ic_role = saved.get("is_ic_role", False)
        parallel_docs = saved.get("parallel_docs_ready", False)

        tracker_entry = {
            "job_id": job_id,
            "company": company,
            "role": role,
            "persona": persona,
            "version": version,
            "date": datetime.now().isoformat(),
            "pre_match_score": saved.get("pre_match_score", 0),
            "post_match_score": post_match_score,
            "score_improvement": post_match_score - saved.get("pre_match_score", 0),
            "ats_coverage_pct": ats_result.get("coverage_pct", 0),
            "audit_score": audit_result.get("score", 0),
            "audit_passed": audit_result.get("passed", False),
            "resume_score_before": baseline_score.get("percentage") if baseline_score else None,
            "resume_score_after": post_score.get("percentage") if post_score else None,
            "resume_grade_before": baseline_score.get("grade") if baseline_score else None,
            "resume_grade_after": post_score.get("grade") if post_score else None,
            "visual_fit_mode": fit_result.get("mode") if fit_result else None,
            "visual_fit_fill": fit_result.get("final_fill") if fit_result else None,
            "certified": certification.get("verdict") if certification else None,
            "certification_blockers": len(certification.get("blockers", [])) if certification else None,
            "certification_warnings": len(certification.get("warnings", [])) if certification else None,
            # Review metadata
            "workflow": "review" if total_sections > 0 else "quick",
            "sections_reviewed": len(approved_sections),
            "total_sections": total_sections,
            "had_user_feedback": bool(user_feedback_text),
            "is_ic_role": is_ic_role,
            "parallel_docs_pregenerated": parallel_docs,
        }
        add_tracker_entry(tracker_entry)

        final = {
            "step": 10,
            "total_steps": 10,
            "label": "Complete",
            "status": "completed",
            "job_id": job_id,
            "company": company,
            "role": role,
            "persona": persona,
            "version": version,
            "phase": "generation",
            "files": files_to_zip + ["Application_Package.zip"],
            "match_score_md": match_content,
            "pre_match_score": saved.get("pre_match_score", 0),
            "post_match_score": post_match_score,
            "audit_result": audit_result,
            "ats_keywords": ats_result,
            "resume_score_before": baseline_score,
            "resume_score_after": post_score,
            "resume_score_delta": score_delta,
            "visual_fit": fit_result,
            "certification": certification,
        }
        yield json.dumps(final)

    except Exception as e:
        yield json.dumps({
            "step": 0,
            "total_steps": 10,
            "status": "error",
            "label": "Generation error",
            "detail": str(e),
        })


# ═══════════════════════════════════════════════════════════════════════════════
# Original full pipeline (backwards compatible)
# ═══════════════════════════════════════════════════════════════════════════════

async def run_pipeline(
    jd_text: str,
    persona: str,
    version: str,
    api_key: str,
) -> AsyncGenerator[str, None]:
    """Run the 12-step resume customization pipeline, yielding SSE events."""

    job_id = str(uuid.uuid4())[:8]
    job_dir = os.path.join(JOBS_DIR, job_id)
    os.makedirs(job_dir, mode=0o755, exist_ok=True)

    try:
        # ── Step 1: Parse inputs ──────────────────────────────────────────
        yield _emit(1, "in_progress")

        if persona == "auto":
            persona = detect_persona(jd_text)

        yield _emit(1, "completed", f"Persona: {persona}, Version: {version}")

        # ── Step 2: Analyze JD ────────────────────────────────────────────
        yield _emit(2, "in_progress")

        jd_analysis = await analyze_jd(api_key, jd_text)

        company = jd_analysis.get("company", "Unknown")
        role = jd_analysis.get("job_title", "Product Manager")
        yield _emit(2, "completed", f"{role} at {company}")

        # ── Step 3: Select and copy template ──────────────────────────────
        yield _emit(3, "in_progress")

        resume_path = os.path.join(job_dir, "Resume.docx")
        copy_template(persona, version, resume_path)

        yield _emit(3, "completed", f"{persona} {version} template")

        # ── Step 4: Read template structure ───────────────────────────────
        yield _emit(4, "in_progress")

        structure = read_template_structure(resume_path)
        bullet_map = get_experience_bullet_map(Document(resume_path))
        structure["experience_entries"] = bullet_map
        yield _emit(4, "completed", f"{structure['total_paragraphs']} paragraphs, {len(structure['sections'])} sections")

        # Score baseline template (before customization)
        baseline_score = None
        try:
            baseline_score = score_resume(resume_path)
        except Exception:
            logger.warning("Score baseline template (before customization) failed (non-blocking)", exc_info=True)

        # ── Step 5: Generate customized content via Claude ────────────────
        yield _emit(5, "in_progress", "Calling Claude API...")

        content = await generate_resume_content(
            api_key, jd_analysis, structure, persona, version, jd_text=jd_text
        )

        if "error" in content:
            yield _emit(5, "warning", f"Partial content: {content.get('error', '')}")
        else:
            yield _emit(5, "completed", "Content generated")

        # ── Step 5b: Run skills library and inject into content ───────────
        # Quick generate doesn't have a review phase, so we explicitly run
        # the skills library to ensure domain-optimized skills selection.
        try:
            skills_result = select_skills(jd_text, version=version)
            content["core_competencies_priority"] = skills_result["core_competencies"]
            content["technical_skills_priority"] = skills_result["technical_skills"]
        except Exception:
            logger.warning("Skills library enrichment failed (non-blocking)", exc_info=True)

        # ── Step 6: Apply edits to resume ─────────────────────────────────
        yield _emit(6, "in_progress")

        doc = Document(resume_path)
        _apply_content_to_resume(doc, content, version, jd_text=jd_text)
        doc.save(resume_path)

        yield _emit(6, "completed", "Resume edits applied")

        # ── Step 7: Quality checks (loop up to 3 times) ──────────────────
        yield _emit(7, "in_progress")

        for attempt in range(3):
            result = verify_resume(resume_path)

            if result["passed"]:
                yield _emit(7, "completed", f"Passed (attempt {attempt + 1})")
                break

            if attempt < 2:
                doc = Document(resume_path)
                bold_sweep(doc)
                doc.save(resume_path)
                yield _emit(7, "in_progress", f"Auto-fixing issues (attempt {attempt + 2})...")
            else:
                detail = f"{result['error_count']} errors, {result['warning_count']} warnings remaining"
                yield _emit(7, "warning", detail)

        # Visual fitting
        try:
            target_pages = 1 if version == "1-Page" else 2
            fit_resume(resume_path, target_pages=target_pages)
        except Exception:
            logger.warning("Visual fitting failed (non-blocking)", exc_info=True)

        # Score customized resume (after edits, with JD for alignment)
        post_score = None
        score_delta = None
        try:
            post_score = score_resume(resume_path, jd_text=jd_text)
            if baseline_score:
                score_delta = compare_scores(baseline_score, post_score)
        except Exception:
            logger.warning("score_delta = None failed (non-blocking)", exc_info=True)

        # ── Step 8: Match Score ───────────────────────────────────────────
        yield _emit(8, "in_progress")

        match_content = await generate_match_score(api_key, jd_text, jd_analysis)
        match_path = os.path.join(job_dir, "Match_Score.docx")
        create_match_score(match_content, company, role, match_path)

        yield _emit(8, "completed")

        # ── Step 9: Cover Letter ──────────────────────────────────────────
        yield _emit(9, "in_progress")

        letter_content = await generate_cover_letter(api_key, jd_text, jd_analysis)
        letter_path = os.path.join(job_dir, "Cover_Letter.docx")
        create_cover_letter(letter_content, company, role, letter_path)

        yield _emit(9, "completed")

        # ── Step 10: Company Brief ────────────────────────────────────────
        yield _emit(10, "in_progress")

        brief_content = await generate_company_brief(api_key, jd_analysis)
        brief_path = os.path.join(job_dir, "Company_Brief.docx")
        create_company_brief(brief_content, company, role, brief_path)

        yield _emit(10, "completed")

        # ── Step 11: Interview Questions ──────────────────────────────────
        yield _emit(11, "in_progress")

        questions_content = await generate_interview_questions(api_key, jd_text, jd_analysis)
        questions_path = os.path.join(job_dir, "Interview_Questions.docx")
        create_interview_questions(questions_content, company, role, questions_path)

        yield _emit(11, "completed")

        # Final certification
        target_pages = 1 if version == "1-Page" else 2
        certification = None
        try:
            certification = certify_resume(
                resume_path,
                target_pages=target_pages,
                jd_text=jd_text,
            )
        except Exception:
            logger.warning("certification = certify_resume( failed (non-blocking)", exc_info=True)

        # ── Step 12: Package deliverables ─────────────────────────────────
        yield _emit(12, "in_progress")

        zip_path = os.path.join(job_dir, "Application_Package.zip")
        files_to_zip = [
            "Resume.docx", "Match_Score.docx", "Cover_Letter.docx",
            "Company_Brief.docx", "Interview_Questions.docx",
        ]
        with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
            for fname in files_to_zip:
                fpath = os.path.join(job_dir, fname)
                if os.path.exists(fpath):
                    zf.write(fpath, fname)

        # ── Application tracker entry (quick generate) ───────────────────
        from datetime import datetime
        tracker_entry = {
            "job_id": job_id,
            "company": company,
            "role": role,
            "persona": persona,
            "version": version,
            "date": datetime.now().isoformat(),
            "resume_score_before": baseline_score.get("percentage") if baseline_score else None,
            "resume_score_after": post_score.get("percentage") if post_score else None,
            "resume_grade_before": baseline_score.get("grade") if baseline_score else None,
            "resume_grade_after": post_score.get("grade") if post_score else None,
            "certified": certification.get("verdict") if certification else None,
            "workflow": "quick",
            "sections_reviewed": 0,
            "total_sections": 0,
            "had_user_feedback": False,
            "is_ic_role": jd_analysis.get("is_ic_role", False),
            "parallel_docs_pregenerated": False,
        }
        add_tracker_entry(tracker_entry)

        final = {
            "step": 12,
            "total_steps": 12,
            "label": "Complete",
            "status": "completed",
            "job_id": job_id,
            "company": company,
            "role": role,
            "persona": persona,
            "version": version,
            "files": files_to_zip + ["Application_Package.zip"],
            "match_score_md": match_content,
            "resume_score_before": baseline_score,
            "resume_score_after": post_score,
            "resume_score_delta": score_delta,
            "certification": certification,
        }
        yield json.dumps(final)

    except Exception as e:
        yield json.dumps({
            "step": 0,
            "status": "error",
            "label": "Pipeline error",
            "detail": str(e),
        })


def get_job_file(job_id: str, filename: str) -> str:
    """Get the path to a specific file in a job's output folder."""
    path = os.path.join(JOBS_DIR, job_id, filename)
    if not os.path.exists(path):
        raise FileNotFoundError(f"File not found: {filename} for job {job_id}")
    return path


def list_job_files(job_id: str) -> List[str]:
    """List all files in a job's output folder."""
    job_dir = os.path.join(JOBS_DIR, job_id)
    if not os.path.exists(job_dir):
        return []
    return [f for f in os.listdir(job_dir) if f.endswith((".docx", ".zip"))]
