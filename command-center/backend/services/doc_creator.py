"""
Document Creator Service — Generates the 4 supporting .docx files.

Creates professionally formatted Word documents for:
1. Match Score & Gap Analysis
2. Cover Letter
3. Company Research Brief
4. Interview Question Bank

Uses python-docx with consistent professional styling.
"""

import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ──────────────────────────────────────────────────────────────────────────────
# Shared Styles
# ──────────────────────────────────────────────────────────────────────────────

FONT_HEADING = "Helvetica"
FONT_BODY = "Helvetica"
COLOR_HEADING = RGBColor(0x1A, 0x1A, 0x2E)  # Dark navy
COLOR_BODY = RGBColor(0x33, 0x33, 0x33)
COLOR_ACCENT = RGBColor(0x2D, 0x6A, 0x4F)  # Green accent


def _setup_doc(title: str, subtitle: str = "") -> Document:
    """Create a new document with consistent header styling."""
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = FONT_BODY
    font.size = Pt(10.5)
    font.color.rgb = COLOR_BODY

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Title
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = COLOR_HEADING
    run.font.name = FONT_HEADING
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Subtitle
    if subtitle:
        p = doc.add_paragraph()
        run = p.add_run(subtitle)
        run.font.size = Pt(11)
        run.font.color.rgb = COLOR_ACCENT
        run.font.name = FONT_BODY
        run.italic = True

    # Divider
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)

    return doc


def _add_heading(doc: Document, text: str, level: int = 2):
    """Add a styled heading."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = FONT_HEADING
    if level == 1:
        run.font.size = Pt(16)
    elif level == 2:
        run.font.size = Pt(13)
    else:
        run.font.size = Pt(11)
    run.font.color.rgb = COLOR_HEADING
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)


def _add_body(doc: Document, text: str):
    """Add body paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = FONT_BODY
    run.font.size = Pt(10.5)
    run.font.color.rgb = COLOR_BODY
    p.paragraph_format.space_after = Pt(6)


def _add_bullet(doc: Document, text: str):
    """Add a bullet point."""
    p = doc.add_paragraph(style="List Bullet")
    # Clear default and add our own
    p.clear()
    run = p.add_run(f"• {text}")
    run.font.name = FONT_BODY
    run.font.size = Pt(10.5)
    run.font.color.rgb = COLOR_BODY
    p.paragraph_format.space_after = Pt(3)


def _markdown_to_docx(doc: Document, markdown_text: str):
    """Convert markdown-formatted text to docx elements."""
    lines = markdown_text.strip().split("\n")
    i = 0

    while i < len(lines):
        line = lines[i].strip()

        if not line:
            i += 1
            continue

        # Heading 1: # Title
        if line.startswith("# ") and not line.startswith("## "):
            _add_heading(doc, line[2:].strip(), level=1)
        # Heading 2: ## Title
        elif line.startswith("## "):
            _add_heading(doc, line[3:].strip(), level=2)
        # Heading 3: ### Title
        elif line.startswith("### "):
            _add_heading(doc, line[4:].strip(), level=3)
        # Bullet point: - text or * text
        elif line.startswith("- ") or line.startswith("* ") or line.startswith("• "):
            bullet_text = line.lstrip("-*• ").strip()
            # Handle bold within bullet: **text**
            bullet_text = re.sub(r"\*\*(.*?)\*\*", r"\1", bullet_text)
            _add_bullet(doc, bullet_text)
        # Numbered list: 1. text
        elif re.match(r"^\d+\.\s", line):
            text = re.sub(r"^\d+\.\s*", "", line).strip()
            text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
            _add_bullet(doc, text)
        # Regular paragraph
        else:
            # Collect consecutive non-empty, non-special lines into one paragraph
            paragraph_lines = [line]
            while i + 1 < len(lines):
                next_line = lines[i + 1].strip()
                if (not next_line or next_line.startswith("#") or
                    next_line.startswith("- ") or next_line.startswith("* ") or
                    next_line.startswith("• ") or re.match(r"^\d+\.\s", next_line)):
                    break
                paragraph_lines.append(next_line)
                i += 1

            full_text = " ".join(paragraph_lines)
            # Strip markdown bold
            full_text = re.sub(r"\*\*(.*?)\*\*", r"\1", full_text)
            _add_body(doc, full_text)

        i += 1


# ──────────────────────────────────────────────────────────────────────────────
# Document Generators
# ──────────────────────────────────────────────────────────────────────────────

def create_match_score(content: str, company: str, role: str, output_path: str):
    """Create the Match Score & Gap Analysis document."""
    doc = _setup_doc(
        "Match Score & Gap Analysis",
        f"{role} at {company}"
    )
    _markdown_to_docx(doc, content)
    doc.save(output_path)


def create_cover_letter(content: str, company: str, role: str, output_path: str):
    """Create the Cover Letter document."""
    doc = _setup_doc("Cover Letter", f"{role} — {company}")

    # Add date
    from datetime import date
    _add_body(doc, date.today().strftime("%B %d, %Y"))

    # Add greeting and body (cover letter is plain text, not markdown)
    paragraphs = content.strip().split("\n\n")
    for para in paragraphs:
        para = para.strip()
        if para:
            _add_body(doc, para)

    # Signature
    doc.add_paragraph()
    _add_body(doc, "Sincerely,")
    p = doc.add_paragraph()
    run = p.add_run("Kiran Gorapalli")
    run.bold = True
    run.font.name = FONT_BODY
    run.font.size = Pt(10.5)

    doc.save(output_path)


def create_company_brief(content: str, company: str, role: str, output_path: str):
    """Create the Company Research Brief document."""
    doc = _setup_doc(
        f"Company Research Brief: {company}",
        f"Prepared for {role} application"
    )
    _markdown_to_docx(doc, content)
    doc.save(output_path)


def create_interview_questions(content: str, company: str, role: str, output_path: str):
    """Create the Interview Question Bank document."""
    doc = _setup_doc(
        "Interview Question Bank",
        f"{role} at {company}"
    )
    _markdown_to_docx(doc, content)
    doc.save(output_path)
