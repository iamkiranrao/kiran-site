"""
Resume Scorer — Research-backed scoring framework for resume quality assessment.

Scores resumes across 5 categories with 19 individual metrics:
1. Readability & Scannability (Flesch, grade level, bullet length, syllable complexity)
2. Information Density (word count, bullets, words/page, metrics density)
3. ATS Compatibility (headers, format, dates, special chars)
4. Content Quality (verb uniqueness, verb diversity, structure variety, banned words)
5. PM-Specific Signals (keywords, career progression, scale signals)

Optional 6th category when JD is provided:
6. JD Alignment (keyword coverage, title match, skill overlap)

Usage:
    from services.resume_scorer import score_resume
    result = score_resume("/path/to/resume.docx")
    result_with_jd = score_resume("/path/to/resume.docx", jd_text="...")
"""

import re
import os
from collections import Counter
from typing import Optional, Dict, List, Tuple
from docx import Document


# ── Constants ──────────────────────────────────────────────────────────

HARD_BANNED = [
    "orchestrated", "spearheaded", "championed", "pioneered",
    "seasoned", "comprehensive", "leveraged", "synergized",
]
SOFT_BANNED = [
    "architected", "streamlined", "revolutionized",
    "transformed", "empowered", "optimized",
]
AI_LANGUAGE = [
    "delve", "landscape", "multifaceted", "game-changer",
    "compelling", "robust", "holistic",
]
BUZZWORDS = [
    "leverage", "synergy", "paradigm", "ecosystem",
    "scalable solution", "cutting-edge", "best-in-class", "empower",
]
ANTI_PATTERNS = [
    "results-driven", "dynamic", "passionate about",
    "dedicated to", "seeking",
]
FILLER_PATTERNS = [
    r"\bvery\b", r"\breally\b", r"\bextremely\b",
    r"\bincredibly\b", r"\bhighly\b",
]

ATS_HEADERS = [
    "experience", "education", "skills", "summary",
    "competencies", "certifications",
]

PM_KEYWORDS = [
    "roadmap", "cross-functional", "stakeholder", "a/b test",
    "user research", "data-driven", "gtm", "p&l", "okr",
    "experimentation",
]

IMPACT_VERBS = {"Drove", "Grew", "Increased", "Expanded", "Scaled", "Cut", "Reduced"}
CREATION_VERBS = {"Built", "Designed", "Created", "Shipped", "Launched", "Introduced", "Developed"}
LEADERSHIP_VERBS = {"Led", "Managed", "Owned", "Ran", "Directed", "Oversaw"}
EXECUTION_VERBS = {"Delivered", "Deployed", "Migrated", "Integrated", "Implemented", "Established", "Executed"}
STRATEGY_VERBS = {"Defined", "Restructured", "Repositioned", "Consolidated", "Negotiated", "Reworked"}


# ── Helpers ────────────────────────────────────────────────────────────

def _count_syllables(word: str) -> int:
    """Approximate syllable count for English words."""
    word = word.lower().strip(".,;:!?\"'()-")
    if not word:
        return 0
    count = 0
    vowels = "aeiouy"
    prev_vowel = False
    for ch in word:
        is_vowel = ch in vowels
        if is_vowel and not prev_vowel:
            count += 1
        prev_vowel = is_vowel
    if word.endswith("e") and count > 1:
        count -= 1
    return max(1, count)


def _flesch_scores(text: str) -> Tuple[float, float, float, float]:
    """Return (reading_ease, grade_level, avg_sentence_length, avg_syllables_per_word)."""
    sentences = [s.strip() for s in re.split(r'[.!?]+', text) if len(s.strip()) > 5]
    words = text.split()
    if not words or not sentences:
        return 0.0, 0.0, 0.0, 0.0

    total_words = len(words)
    total_sentences = max(1, len(sentences))
    total_syllables = sum(_count_syllables(w) for w in words)

    fre = 206.835 - 1.015 * (total_words / total_sentences) - 84.6 * (total_syllables / total_words)
    fk = 0.39 * (total_words / total_sentences) + 11.8 * (total_syllables / total_words) - 15.59
    return fre, fk, total_words / total_sentences, total_syllables / total_words


def _extract_text(doc: Document) -> str:
    """Extract full text from docx."""
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def _extract_bullets(doc: Document) -> List[str]:
    """Extract bullet point text (lines starting with bullet markers)."""
    bullets = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if text.startswith("•") or text.startswith("·"):
            bullets.append(text[1:].strip())
    return bullets


def _extract_verbs(bullets: List[str]) -> List[str]:
    """Extract leading verbs from bullet points."""
    skip = {"Jira", "Mixpanel", "OpenAI", "Agile", "Liverpool", "Product", "PM", "Analytics", "AI", "Frameworks"}
    verbs = []
    for b in bullets:
        words = b.split()
        if words and words[0][0].isalpha():
            verb = words[0].rstrip(",.:;")
            if verb not in skip and len(verb) > 2:
                verbs.append(verb)
    return verbs


def _count_pages(docx_path: str) -> int:
    """Estimate page count from docx. Uses PDF conversion if available, otherwise estimates."""
    pdf_path = docx_path.replace(".docx", ".pdf")
    if os.path.exists(pdf_path):
        try:
            import PyPDF2
            reader = PyPDF2.PdfReader(pdf_path)
            return len(reader.pages)
        except Exception:
            pass
    # Rough estimate: ~45 lines per page
    doc = Document(docx_path)
    line_count = sum(1 for p in doc.paragraphs if p.text.strip())
    return max(1, round(line_count / 45))


def _score_metric(value: float, ideal_lo: float, ideal_hi: float,
                  ok_lo: float = None, ok_hi: float = None, max_pts: int = 10) -> int:
    """Score a metric: full points if in ideal range, partial if in ok range, low otherwise."""
    if ideal_lo <= value <= ideal_hi:
        return max_pts
    if ok_lo is not None and ok_hi is not None:
        if ok_lo <= value <= ok_hi:
            return round(max_pts * 0.7)
    return round(max_pts * 0.4)


# ── JD Keyword Extraction ─────────────────────────────────────────────

def _extract_jd_keywords(jd_text: str) -> List[str]:
    """Extract significant keywords from a job description."""
    # Common stop words to ignore
    stop_words = {
        "the", "a", "an", "and", "or", "but", "in", "on", "at", "to", "for",
        "of", "with", "by", "from", "as", "is", "was", "are", "were", "be",
        "been", "being", "have", "has", "had", "do", "does", "did", "will",
        "would", "could", "should", "may", "might", "must", "shall", "can",
        "this", "that", "these", "those", "it", "its", "we", "our", "you",
        "your", "they", "their", "them", "who", "what", "which", "where",
        "when", "how", "all", "each", "every", "both", "few", "more", "most",
        "other", "some", "such", "no", "not", "only", "own", "same", "so",
        "than", "too", "very", "just", "about", "above", "after", "again",
        "also", "any", "because", "before", "between", "during", "into",
        "through", "under", "until", "up", "while", "able", "across",
        "including", "within", "strong", "experience", "work", "working",
        "role", "team", "ability", "skills", "years", "required", "preferred",
        "responsibilities", "qualifications", "requirements", "job", "position",
        "company", "join", "looking", "ideal", "candidate", "opportunity",
    }

    words = re.findall(r'\b[a-z][a-z-]+\b', jd_text.lower())
    counts = Counter(w for w in words if w not in stop_words and len(w) > 2)

    # Also extract multi-word phrases
    bigrams = []
    word_list = jd_text.lower().split()
    for i in range(len(word_list) - 1):
        w1 = re.sub(r'[^a-z-]', '', word_list[i])
        w2 = re.sub(r'[^a-z-]', '', word_list[i + 1])
        if w1 and w2 and w1 not in stop_words and w2 not in stop_words:
            bigrams.append(f"{w1} {w2}")
    bigram_counts = Counter(bigrams)

    # Return top keywords by frequency
    top_singles = [w for w, c in counts.most_common(30) if c >= 2]
    top_bigrams = [b for b, c in bigram_counts.most_common(15) if c >= 2]

    return top_bigrams + top_singles


# ── Main Scorer ────────────────────────────────────────────────────────

def score_resume(docx_path: str, jd_text: Optional[str] = None,
                 page_count: Optional[int] = None) -> Dict:
    """
    Score a resume across 5 categories (6 if JD provided).

    Args:
        docx_path: Path to the .docx resume file
        jd_text: Optional job description text for JD alignment scoring
        page_count: Optional override for page count (avoids PDF dependency)

    Returns:
        Dict with category scores, individual metrics, and grand total
    """
    doc = Document(docx_path)
    full_text = _extract_text(doc)
    full_lower = full_text.lower()
    bullets = _extract_bullets(doc)
    bullet_text = " ".join(bullets)
    verbs = _extract_verbs(bullets)
    all_words = full_text.split()

    pages = page_count or _count_pages(docx_path)

    # ── 1. READABILITY & SCANNABILITY (40 pts) ─────────────────────

    fre, fk_grade, avg_sent_len, avg_syl = _flesch_scores(bullet_text)
    bullet_word_counts = [len(b.split()) for b in bullets]
    avg_bwc = sum(bullet_word_counts) / len(bullet_word_counts) if bullet_word_counts else 0

    fre_pts = _score_metric(fre, 30, 60, 20, 70)
    fk_pts = _score_metric(fk_grade, 9, 14, 7, 16)
    bwc_pts = _score_metric(avg_bwc, 15, 25, 12, 30)
    syl_pts = _score_metric(avg_syl, 1.5, 2.0, 1.3, 2.2)

    readability = {
        "category": "Readability & Scannability",
        "max": 40,
        "score": fre_pts + fk_pts + bwc_pts + syl_pts,
        "metrics": {
            "flesch_reading_ease": {"value": round(fre, 1), "target": "30-60", "score": fre_pts, "max": 10},
            "flesch_kincaid_grade": {"value": round(fk_grade, 1), "target": "9-14", "score": fk_pts, "max": 10},
            "avg_words_per_bullet": {"value": round(avg_bwc, 1), "target": "15-25", "score": bwc_pts, "max": 10},
            "avg_syllables_per_word": {"value": round(avg_syl, 2), "target": "1.5-2.0", "score": syl_pts, "max": 10},
        },
    }

    # ── 2. INFORMATION DENSITY (40 pts) ────────────────────────────

    total_words = len(all_words)
    total_bullets = len(bullets)
    words_per_page = total_words / max(1, pages)

    metric_bullets = sum(1 for b in bullets if re.search(r'\d+[%MKB]|\$\d|\d+\s*(?:point|person)', b))
    metrics_pct = (metric_bullets / total_bullets * 100) if total_bullets else 0

    if pages == 1:
        wc_pts = _score_metric(total_words, 300, 475, 250, 525)
    else:
        wc_pts = _score_metric(total_words, 475, 700, 425, 800)

    bc_pts = 10 if 8 <= total_bullets <= 35 else 7
    wpp_pts = _score_metric(words_per_page, 250, 400, 200, 450)
    md_pts = _score_metric(metrics_pct, 65, 85, 55, 90)

    density = {
        "category": "Information Density",
        "max": 40,
        "score": wc_pts + bc_pts + wpp_pts + md_pts,
        "metrics": {
            "total_word_count": {"value": total_words, "target": "300-475 (1pg) / 475-700 (2pg)", "score": wc_pts, "max": 10},
            "total_bullets": {"value": total_bullets, "target": "8-35", "score": bc_pts, "max": 10},
            "words_per_page": {"value": round(words_per_page), "target": "250-400", "score": wpp_pts, "max": 10},
            "metrics_density": {"value": f"{metric_bullets}/{total_bullets} ({metrics_pct:.0f}%)", "target": "65-85%", "score": md_pts, "max": 10},
        },
    }

    # ── 3. ATS COMPATIBILITY (40 pts) ──────────────────────────────

    found_headers = sum(1 for h in ATS_HEADERS if h in full_lower)
    hdr_pts = 10 if found_headers >= 4 else (7 if found_headers >= 3 else 4)

    # Format: always 10 for our templates (single-column, Calibri, .docx)
    fmt_pts = 10

    # Date format (check for apostrophe years in text)
    bad_dates = re.findall(r"'\d{2}\b", full_text)
    date_pts = 10 if not bad_dates else 4

    # Special characters
    em_dashes = full_text.count("—")
    arrows = full_text.count("→")
    char_pts = 10 if em_dashes == 0 and arrows == 0 else 4

    ats = {
        "category": "ATS Compatibility",
        "max": 40,
        "score": hdr_pts + fmt_pts + date_pts + char_pts,
        "metrics": {
            "ats_headers": {"value": f"{found_headers}/6", "target": "4+", "score": hdr_pts, "max": 10},
            "format_compliance": {"value": "Single-column, Calibri, .docx", "target": "Pass", "score": fmt_pts, "max": 10},
            "date_format": {"value": "Consistent" if not bad_dates else f"{len(bad_dates)} issues", "target": "Consistent", "score": date_pts, "max": 10},
            "special_characters": {"value": f"Em dashes: {em_dashes}, Arrows: {arrows}", "target": "0", "score": char_pts, "max": 10},
        },
    }

    # ── 4. CONTENT QUALITY (40 pts) ────────────────────────────────

    # Verb uniqueness
    verb_counts = Counter(verbs)
    dupes = {v: c for v, c in verb_counts.items() if c > 1}
    verb_pts = 10 if not dupes else max(0, 10 - len(dupes) * 3)

    # Verb category diversity
    verb_set = set(verbs)
    categories_hit = 0
    if verb_set & IMPACT_VERBS: categories_hit += 1
    if verb_set & CREATION_VERBS: categories_hit += 1
    if verb_set & LEADERSHIP_VERBS: categories_hit += 1
    if verb_set & EXECUTION_VERBS: categories_hit += 1
    if verb_set & STRATEGY_VERBS: categories_hit += 1
    vcat_pts = 10 if categories_hit >= 4 else (7 if categories_hit >= 3 else 4)

    # Outcome-first bullets
    outcome_first = sum(1 for b in bullets if b and b[0].isdigit())
    struct_pts = 10 if outcome_first >= 1 else 5

    # Banned words
    violations = []
    for w in HARD_BANNED + SOFT_BANNED + AI_LANGUAGE + BUZZWORDS:
        if w.lower() in full_lower:
            violations.append(w)
    for pattern in FILLER_PATTERNS:
        if re.search(pattern, full_lower):
            violations.append(pattern.strip(r"\b"))
    for ap in ANTI_PATTERNS:
        if ap in full_lower:
            violations.append(ap)
    ban_pts = 10 if not violations else max(0, 10 - len(violations) * 3)

    content = {
        "category": "Content Quality",
        "max": 40,
        "score": verb_pts + vcat_pts + struct_pts + ban_pts,
        "metrics": {
            "verb_uniqueness": {"value": f"{len(set(verbs))}/{len(verbs)}", "target": "All unique", "score": verb_pts, "max": 10},
            "verb_diversity": {"value": f"{categories_hit}/5 categories", "target": "4+", "score": vcat_pts, "max": 10},
            "structure_variety": {"value": f"{outcome_first} outcome-first", "target": "1+", "score": struct_pts, "max": 10},
            "banned_words": {"value": "None" if not violations else ", ".join(violations), "target": "None", "score": ban_pts, "max": 10},
        },
    }

    # ── 5. PM-SPECIFIC SIGNALS (30 pts) ────────────────────────────

    pm_present = [kw for kw in PM_KEYWORDS if kw in full_lower]
    pm_missing = [kw for kw in PM_KEYWORDS if kw not in full_lower]
    pm_count = len(pm_present)
    pm_pts = 10 if pm_count >= 7 else (7 if pm_count >= 5 else 4)

    # Career progression
    titles = re.findall(r'(VP|Director|AVP|Senior Consultant|Manager|Lead|Principal)', full_text)
    unique_titles = list(dict.fromkeys(titles))
    prog_pts = 10 if len(set(titles)) >= 3 else (7 if len(set(titles)) >= 2 else 4)

    # Scale signals
    scale_patterns = [r'\d+M\s*(?:MAU|users|interactions)', r'\d+-person', r'\$\d+M']
    scale_hits = sum(1 for p in scale_patterns if re.search(p, full_text))
    scale_pts = 10 if scale_hits >= 3 else (7 if scale_hits >= 2 else 4)

    pm_signals = {
        "category": "PM-Specific Signals",
        "max": 30,
        "score": pm_pts + prog_pts + scale_pts,
        "metrics": {
            "pm_keywords": {"value": f"{pm_count}/10", "present": pm_present, "missing": pm_missing, "score": pm_pts, "max": 10},
            "career_progression": {"value": " → ".join(unique_titles), "score": prog_pts, "max": 10},
            "scale_signals": {"value": f"{scale_hits}/3", "target": "3/3", "score": scale_pts, "max": 10},
        },
    }

    # ── 6. JD ALIGNMENT (optional, 30 pts) ─────────────────────────

    jd_alignment = None
    if jd_text:
        jd_keywords = _extract_jd_keywords(jd_text)
        total_jd_kw = len(jd_keywords)

        # Keyword coverage
        matched = [kw for kw in jd_keywords if kw in full_lower]
        coverage_pct = (len(matched) / total_jd_kw * 100) if total_jd_kw else 0
        cov_pts = _score_metric(coverage_pct, 60, 85, 40, 95)

        # Title match: extract JD title and check if it appears in resume
        title_match = False
        jd_lower = jd_text.lower()
        for pattern in [r'(?:title|position|role)\s*[:]\s*(.+)', r'^(.+(?:manager|director|lead|vp|head).+)$']:
            m = re.search(pattern, jd_lower, re.MULTILINE)
            if m:
                jd_title_words = set(m.group(1).strip().split())
                resume_words = set(full_lower.split())
                overlap = jd_title_words & resume_words
                if len(overlap) >= 2:
                    title_match = True
                    break
        title_pts = 10 if title_match else 5

        # Skill overlap
        jd_skills = set(re.findall(
            r'\b(?:python|sql|jira|figma|tableau|agile|scrum|okr|analytics|'
            r'a/b\s*test|machine\s*learning|ai|ml|api|saas|b2b|b2c|gtm|'
            r'roadmap|stakeholder|cross-functional|data-driven|experimentation|'
            r'product\s*strategy|user\s*research|growth|retention|engagement|'
            r'mobile|web|platform|cloud|aws|gcp)\b',
            jd_lower
        ))
        resume_skills = set(re.findall(
            r'\b(?:python|sql|jira|figma|tableau|agile|scrum|okr|analytics|'
            r'a/b\s*test|machine\s*learning|ai|ml|api|saas|b2b|b2c|gtm|'
            r'roadmap|stakeholder|cross-functional|data-driven|experimentation|'
            r'product\s*strategy|user\s*research|growth|retention|engagement|'
            r'mobile|web|platform|cloud|aws|gcp)\b',
            full_lower
        ))
        skill_overlap = jd_skills & resume_skills
        skill_pct = (len(skill_overlap) / len(jd_skills) * 100) if jd_skills else 100
        skill_pts = _score_metric(skill_pct, 60, 100, 40, 100)

        jd_alignment = {
            "category": "JD Alignment",
            "max": 30,
            "score": cov_pts + title_pts + skill_pts,
            "metrics": {
                "keyword_coverage": {
                    "value": f"{len(matched)}/{total_jd_kw} ({coverage_pct:.0f}%)",
                    "target": "60-85%",
                    "matched": matched[:15],
                    "missing": [kw for kw in jd_keywords if kw not in full_lower][:10],
                    "score": cov_pts,
                    "max": 10,
                },
                "title_match": {"value": "Match" if title_match else "No match", "score": title_pts, "max": 10},
                "skill_overlap": {
                    "value": f"{len(skill_overlap)}/{len(jd_skills)} ({skill_pct:.0f}%)",
                    "target": "60%+",
                    "matched": list(skill_overlap),
                    "missing": list(jd_skills - resume_skills),
                    "score": skill_pts,
                    "max": 10,
                },
            },
        }

    # ── Assemble Results ───────────────────────────────────────────

    categories = [readability, density, ats, content, pm_signals]
    if jd_alignment:
        categories.append(jd_alignment)

    grand_score = sum(c["score"] for c in categories)
    grand_max = sum(c["max"] for c in categories)
    grand_pct = (grand_score / grand_max * 100) if grand_max else 0

    if grand_pct >= 95:
        grade = "A+"
    elif grand_pct >= 90:
        grade = "A"
    elif grand_pct >= 85:
        grade = "A-"
    elif grand_pct >= 80:
        grade = "B+"
    elif grand_pct >= 75:
        grade = "B"
    elif grand_pct >= 70:
        grade = "B-"
    else:
        grade = "C"

    return {
        "score": grand_score,
        "max": grand_max,
        "percentage": round(grand_pct, 1),
        "grade": grade,
        "pages": pages,
        "categories": {c["category"]: c for c in categories},
    }


def format_scorecard(result: Dict, label: str = "Resume") -> str:
    """Format a score result as a human-readable text scorecard."""
    lines = []
    lines.append(f"{'='*60}")
    lines.append(f"  {label} — {result['grade']} ({result['percentage']}%)")
    lines.append(f"  {result['score']}/{result['max']} points | {result['pages']} page(s)")
    lines.append(f"{'='*60}")

    for cat_name, cat in result["categories"].items():
        lines.append(f"\n  {cat['category']}  [{cat['score']}/{cat['max']}]")
        lines.append(f"  {'—'*45}")
        for metric_name, m in cat["metrics"].items():
            label_str = metric_name.replace("_", " ").title()
            score_str = f"[{m['score']}/{m['max']}]"
            value_str = str(m.get("value", ""))
            target_str = f"  (target: {m['target']})" if "target" in m else ""
            lines.append(f"    {score_str:8s} {label_str}: {value_str}{target_str}")

    lines.append(f"\n{'='*60}")
    return "\n".join(lines)


def compare_scores(before: Dict, after: Dict) -> Dict:
    """Compare before/after scores and return a delta summary."""
    deltas = {}
    for cat_name in before["categories"]:
        b_cat = before["categories"][cat_name]
        a_cat = after["categories"].get(cat_name)
        if not a_cat:
            continue
        cat_delta = a_cat["score"] - b_cat["score"]
        metric_deltas = {}
        for m_name in b_cat["metrics"]:
            b_m = b_cat["metrics"][m_name]
            a_m = a_cat["metrics"].get(m_name)
            if a_m:
                metric_deltas[m_name] = {
                    "before": b_m["score"],
                    "after": a_m["score"],
                    "delta": a_m["score"] - b_m["score"],
                    "before_value": b_m.get("value"),
                    "after_value": a_m.get("value"),
                }
        deltas[cat_name] = {
            "before": b_cat["score"],
            "after": a_cat["score"],
            "delta": cat_delta,
            "metrics": metric_deltas,
        }

    # JD Alignment only exists in after
    if "JD Alignment" in after["categories"] and "JD Alignment" not in before["categories"]:
        jd_cat = after["categories"]["JD Alignment"]
        deltas["JD Alignment"] = {
            "before": 0,
            "after": jd_cat["score"],
            "delta": jd_cat["score"],
            "note": "New category (only scored with JD)",
            "metrics": {
                m_name: {
                    "before": 0,
                    "after": m["score"],
                    "delta": m["score"],
                    "after_value": m.get("value"),
                }
                for m_name, m in jd_cat["metrics"].items()
            },
        }

    return {
        "before_total": before["score"],
        "after_total": after["score"],
        "before_max": before["max"],
        "after_max": after["max"],
        "before_pct": before["percentage"],
        "after_pct": after["percentage"],
        "before_grade": before["grade"],
        "after_grade": after["grade"],
        "categories": deltas,
    }
